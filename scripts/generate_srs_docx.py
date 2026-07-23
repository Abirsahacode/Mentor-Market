from pathlib import Path
import sys

sys.path.insert(0, "/tmp/mentor_market_docx")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "documentation"
OUTPUT = OUTPUT_DIR / "Mentor_Market_SRS.docx"

NAVY = "17324D"
TEAL = "0C7C86"
PALE = "E8F3F4"
LIGHT = "F3F6F8"
GRAY = "5C6873"
WHITE = "FFFFFF"
SERIF_FONT = "Times New Roman"
MONO_FONT = "Courier New"


def apply_font(run, name=SERIF_FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def apply_style_font(style, name, size, color=None, bold=None, italic=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style._element.rPr.rFonts.set(qn("w:cs"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    apply_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, WHITE, 9)
        shade(table.rows[0].cells[i], NAVY)
        table.rows[0].cells[i].paragraphs[0].paragraph_format.keep_with_next = True
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_props.append(OxmlElement("w:tblHeader"))
    header_props.append(OxmlElement("w:cantSplit"))
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        cant_split = OxmlElement("w:cantSplit")
        table.rows[-1]._tr.get_or_add_trPr().append(cant_split)
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, None, 8.5)
            if row_idx % 2:
                shade(cells[i], LIGHT)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_req(doc, req_id, title, statements, priority="High"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{req_id} — {title}")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    meta = p.add_run(f"  [{priority}]")
    meta.italic = True
    meta.font.size = Pt(9)
    meta.font.color.rgb = RGBColor.from_string(GRAY)
    for statement in statements:
        add_bullet(doc, statement, 1)


def add_field(paragraph, instruction, display_text=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_static_toc(doc, entries):
    """Add a visible TOC that also acts as the cached result of a Word TOC field."""
    paragraphs = []
    for title, level, page in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.28 * (level - 1))
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.45), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        r = p.add_run(title)
        apply_font(r, size=11, bold=level == 1, color=NAVY if level == 1 else None)
        page_run = p.add_run(f"\t{page}")
        apply_font(page_run, size=11, bold=level == 1)
        paragraphs.append(p)

    first_run = paragraphs[0].runs[0]._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    first_run.insert(0, separate)
    first_run.insert(0, instr)
    first_run.insert(0, begin)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraphs[-1].runs[-1]._r.append(end)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    styles = doc.styles
    normal = styles["Normal"]
    apply_style_font(normal, SERIF_FONT, 11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size, color in [
        ("Title", 26, NAVY),
        ("Subtitle", 15, TEAL),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 12, NAVY),
    ]:
        style = styles[style_name]
        apply_style_font(
            style,
            SERIF_FONT,
            size,
            color=color,
            bold=style_name != "Subtitle",
            italic=style_name == "Subtitle",
        )
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 1"].paragraph_format.page_break_before = False
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    styles["Heading 3"].paragraph_format.keep_with_next = True


def add_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = "MENTOR MARKET  |  SOFTWARE REQUIREMENTS SPECIFICATION"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        apply_font(run, size=8, bold=True, color=TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Mentor Market  •  Version 1.0  •  ")
    add_field(footer, "PAGE", "1")
    for run in footer.runs:
        apply_font(run, size=8, color=GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MENTOR")
    apply_font(run, size=28, bold=True, italic=True, color=NAVY)
    run = p.add_run(" MARKET")
    apply_font(run, size=28, bold=True, italic=True, color=TEAL)
    p = doc.add_paragraph("A Full-Stack EdTech Tutoring Marketplace", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph("Software Requirements Specification", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("SRS", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(62)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    details = [
        ("Document version", "1.0"),
        ("Prepared for", "Mentor Market Project"),
        ("Prepared by", "Mentor Market Project Team"),
        ("Date", "24 July 2026"),
    ]
    for i, (label, value) in enumerate(details):
        set_cell_text(table.cell(i, 0), label, True, WHITE, 10)
        shade(table.cell(i, 0), NAVY)
        set_cell_text(table.cell(i, 1), value, False, None, 10)
        if i % 2:
            shade(table.cell(i, 1), LIGHT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Students find mentors. Mentors find learners. Learning relationships grow.")
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_front_matter_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    apply_font(run, size=17, bold=True, color=NAVY)
    return p


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(content_section)

    add_front_matter_heading(doc, "Document Control")
    add_table(doc, ["Field", "Value"], [
        ("Document title", "Software Requirements Specification — Mentor Market"),
        ("Version", "1.0"),
        ("Status", "Baseline"),
        ("Intended audience", "Project team, course instructor, testers, maintainers, and evaluators"),
        ("Source baseline", "Mentor Market repository implementation and database schema as inspected on 24 July 2026"),
    ], [2, 4.8])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    apply_font(p.add_run("Revision History"), size=14, bold=True, color=TEAL)
    add_table(doc, ["Version", "Date", "Author", "Description"], [
        ("1.0", "24 Jul 2026", "Mentor Market Project Team", "Initial project-specific SRS baseline"),
    ], [0.8, 1.1, 2.0, 3.0])
    doc.add_page_break()
    add_front_matter_heading(doc, "Table of Contents")
    add_static_toc(doc, [
        ("1. Introduction", 1, 5),
        ("1.1 Purpose", 2, 5),
        ("1.2 Scope", 2, 5),
        ("1.3 Definitions, Acronyms, and Abbreviations", 2, 6),
        ("1.4 References", 2, 6),
        ("1.5 Overview", 2, 7),
        ("2. Overall Description", 1, 7),
        ("2.1 Product Perspective", 2, 7),
        ("2.2 Major Product Features", 2, 7),
        ("2.3 User Classes and Characteristics", 2, 8),
        ("2.4 Operating Environment", 2, 9),
        ("2.5 Constraints", 2, 9),
        ("2.6 Assumptions and Dependencies", 2, 10),
        ("3. System Requirements", 1, 10),
        ("3.1 Functional Requirements", 2, 10),
        ("3.1.1 Authentication and Profiles", 3, 10),
        ("3.1.2 Discovery and Marketplace", 3, 11),
        ("3.1.3 Applications, Bookings, and Communication", 3, 12),
        ("3.1.4 Learning Management", 3, 13),
        ("3.1.5 Finance", 3, 14),
        ("3.1.6 Trust, Safety, and Administration", 3, 15),
        ("3.1.7 Business Rules", 3, 16),
        ("3.2 Non-Functional Requirements", 2, 17),
        ("3.2.1 Performance Requirements", 3, 17),
        ("3.2.2 Security and Privacy Requirements", 3, 17),
        ("3.2.3 Reliability and Availability Requirements", 3, 18),
        ("3.2.4 Usability and Accessibility Requirements", 3, 18),
        ("3.2.5 Maintainability, Testability, Portability, and Scalability", 3, 18),
        ("3.3 External Interface Requirements", 2, 19),
        ("3.3.1 User Interfaces", 3, 19),
        ("3.3.2 Hardware Interfaces", 3, 19),
        ("3.3.3 Software Interfaces", 3, 20),
        ("3.3.4 Communication Interfaces", 3, 20),
        ("4. Technology Stack & Architectural Overview", 1, 21),
        ("4.1 Technology Stack Components", 2, 21),
        ("4.2 High-Level Architecture", 2, 21),
        ("4.3 Data Model and Integrity", 2, 22),
        ("4.3.1 Core Entities", 3, 22),
        ("4.3.2 Relationship Model", 3, 23),
        ("4.3.3 Data Integrity and Retention", 3, 23),
        ("4.4 Key Use Cases", 2, 24),
        ("4.4.1 Booking State Model", 3, 24),
        ("5. Tentative Development Plan (Agile Methodology)", 1, 25),
        ("5.1 Definition of Done", 2, 26),
        ("6. Acceptance Criteria", 1, 26),
        ("6.1 System Acceptance Criteria", 2, 26),
        ("6.2 Requirement Traceability Matrix", 2, 27),
        ("7. Conclusion", 1, 28),
        ("Appendix A — Out of Scope for the MVP", 1, 29),
        ("Appendix B — Risks and Mitigations", 1, 30),
    ])
    note = doc.add_paragraph(
        "This contents table is embedded for immediate viewing and is configured to refresh automatically in Microsoft Word."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    apply_font(note.runs[0], size=9, italic=True, color=GRAY)
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This Software Requirements Specification defines the functional and non-functional requirements for Mentor Market, "
        "a full-stack web marketplace that connects students with tutors and supports the complete learning relationship—from "
        "discovery and hiring to classes, learning resources, assessment, payment, progress tracking, and review. It provides "
        "a shared baseline for design, implementation, testing, evaluation, and future maintenance."
    )
    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph("Mentor Market serves three primary roles:")
    add_bullet(doc, "Students discover and compare tutor services, publish tutoring requests, hire tutors, book classes, communicate, complete learning activities, make simulated payments, track progress, and review tutors.")
    add_bullet(doc, "Tutors publish services, apply to student requests, manage bookings, communicate, distribute materials, create assignments and quizzes, review earnings, request withdrawals, and submit credentials for verification.")
    add_bullet(doc, "Administrators oversee users, tutor verification, marketplace records, payments, withdrawals, reports, and platform analytics.")
    doc.add_paragraph(
        "The product is a responsive React web application backed by an Express REST API and a normalized MySQL database. "
        "The academic MVP stores file references as URLs and simulates payments; it does not process real money or upload files to third-party storage."
    )
    doc.add_heading("1.3 Definitions, Acronyms, and Abbreviations", level=2)
    add_table(doc, ["Term", "Meaning"], [
        ("API", "Application Programming Interface; the REST communication layer between frontend and backend."),
        ("JWT", "JSON Web Token used for authenticated sessions and role claims."),
        ("RBAC", "Role-Based Access Control restricting operations to student, tutor, or administrator roles."),
        ("MVP", "Minimum Viable Product; the implemented academic product baseline."),
        ("REST", "Representational State Transfer style used by the HTTP API."),
        ("CRUD", "Create, Read, Update, and Delete operations."),
        ("Trial class", "A limited introductory booking offered by a tutor service."),
        ("Verification", "Administrator review of a tutor's credentials and supporting evidence."),
        ("Platform commission", "Ten percent of a paid booking amount retained in the simulated ledger."),
    ], [1.4, 5.4])
    doc.add_heading("1.4 References", level=2)
    for text in [
        "IEEE/ISO/IEC 29148 principles for requirements engineering and SRS structure.",
        "Mentor Market README, frontend routes, backend REST routes, tests, and MySQL schema.",
        "React, Express.js, Node.js, MySQL, JWT, bcryptjs, Axios, Vite, Helmet, and express-validator documentation.",
        "The supplied WeHeal CSE470 SRS document, used as a structural and academic presentation reference.",
    ]:
        add_bullet(doc, text)
    doc.add_heading("1.5 Overview", level=2)
    doc.add_paragraph(
        "Section 2 describes the product, users, environment, constraints, and assumptions. Section 3 defines functional, "
        "non-functional, and interface requirements. Section 4 describes the technology stack, architecture, data model, and "
        "key use cases. Section 5 presents the agile development plan, Section 6 provides acceptance and traceability criteria, "
        "and Section 7 concludes the specification."
    )

    doc.add_heading("2. Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph(
        "Mentor Market is a standalone, three-tier web application. A React single-page application communicates with an "
        "Express.js MVC API over HTTP/JSON. The API applies authentication, role checks, validation, rate limiting, and business "
        "rules before reading or writing MySQL. The platform combines a two-sided service marketplace with learning-management, "
        "communication, trust, moderation, and simulated financial capabilities."
    )
    doc.add_heading("2.2 Major Product Features", level=2)
    features = [
        ("Identity and access", "Student/tutor registration, login for all provisioned roles, JWT sessions, and protected role routes."),
        ("Tutor discovery", "Search, filters, profile details, verification badges, ratings, saved tutors, saved courses, recent views, recommendations, and comparison."),
        ("Two-sided marketplace", "Tutor service listings, student requests, tutor proposals, application decisions, and bookings."),
        ("Learning workspace", "Messages, study materials, assignments, submissions, grading, quizzes, scoring, and progress analytics."),
        ("Finance", "Mock payment ledger, 10% commission, tutor earnings, and withdrawal requests."),
        ("Trust and safety", "Booking-linked reviews, tutor credential verification, message reporting, safety reports, user suspension, and administrative moderation."),
        ("Administration", "Role-protected dashboard, aggregate statistics, and marketplace resource management."),
    ]
    add_table(doc, ["Feature group", "Summary"], features, [1.7, 5.1])
    doc.add_heading("2.3 User Classes and Characteristics", level=2)
    add_table(doc, ["User class", "Characteristics and goals", "Access level"], [
        ("Visitor", "Browses the home page, tutor catalog, tutor details, student requests, and informational pages.", "Public, read-only"),
        ("Student", "Seeks tutoring, compares choices, books and pays, communicates, completes learning work, and tracks progress.", "Authenticated student features"),
        ("Tutor", "Markets services, responds to demand, teaches students, evaluates work, and manages earnings.", "Authenticated tutor features"),
        ("Administrator", "Maintains trust, safety, quality, user status, verification, and marketplace oversight.", "Protected platform-wide management"),
    ], [1.1, 4.0, 1.7])
    doc.add_heading("2.4 Operating Environment", level=2)
    for item in [
        "Client: modern versions of Chrome, Firefox, Edge, and Safari on desktop, tablet, and mobile.",
        "Frontend: React 19, React Router, Axios, and Vite; responsive CSS with the Inter variable font.",
        "Server: Node.js with Express.js, deployable on a Linux-compatible host.",
        "Database: MySQL 8 with mysql2 connection pooling and utf8mb4 character encoding.",
        "Transport: JSON over HTTP/HTTPS between frontend and API.",
        "Development and verification: npm, Node's test runner, Supertest, and Postman.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("2.5 Constraints", level=2)
    for item in [
        "The data model is relational and must remain compatible with MySQL 8.",
        "Authentication uses JWT and bcrypt-compatible password hashing.",
        "Authorization is restricted to the fixed roles student, tutor, and admin.",
        "Payments are simulated; no production gateway, settlement, or refund transfer is in scope.",
        "Material, proof, image, and video files are represented by URLs; managed binary upload is outside the MVP.",
        "A booking must reference a valid student and tutor; schedule conflicts must be prevented for active time slots.",
        "A review must be linked to a booking, have a rating from 1 to 5, and be unique per booking/reviewer pair.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("2.6 Assumptions and Dependencies", level=2)
    for item in [
        "Users have a stable internet connection and a valid email address.",
        "Tutors accurately enter qualifications, availability, prices, locations, and supporting URLs.",
        "Administrators manually assess submitted tutor credentials and reports.",
        "The configured MySQL service, JWT secret, API base URL, and CORS origins are available in the environment.",
        "Seed data may be used for demonstration, but production-like credentials must not be embedded in source code.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. System Requirements", level=1)
    doc.add_heading("3.1 Functional Requirements", level=2)
    requirements = [
        ("FR-AUTH-01", "Registration", [
            "The system shall self-register a student or tutor with a full name, unique valid email, and password of at least eight characters; administrator accounts shall be provisioned separately.",
            "The system shall hash the password before persistence and create the corresponding student or tutor profile where applicable.",
            "The system shall reject invalid, incomplete, duplicate, or unsupported registration data."
        ]),
        ("FR-AUTH-02", "Login, session, and logout", [
            "The system shall authenticate an active user by email and password and issue a JWT containing identity and role.",
            "The system shall expose the current authenticated user and shall reject invalid or expired credentials.",
            "The system shall provide logout behavior that clears the client session."
        ]),
        ("FR-PROF-01", "Profile management", [
            "Students shall update personal and learning-profile details including institution, class level, subjects, location, goals, and biography.",
            "Tutors shall update qualifications, experience, subjects, teaching mode, hourly rate, location, availability, and biography.",
            "Users shall update permitted account information without changing protected role or moderation fields."
        ]),
        ("FR-DISC-01", "Tutor discovery and filtering", [
            "Visitors and students shall browse active, sufficiently complete tutor profiles and active services and filter by subject, mode, location, price, rating, experience, availability, verification, and trial support where available.",
            "The system shall debounce interactive text search and paginate or progressively load results.",
            "The system shall show tutor profile, service, rating, verification, and availability details."
        ]),
        ("FR-DISC-02", "Saved and personalized discovery", [
            "Authenticated students shall save and remove tutors and courses.",
            "The system shall record recent course views and return recent history.",
            "The student discovery experience shall use account engagement and browser-local subject interests for ordering and shall allow comparison of two or three courses."
        ], "Medium"),
        ("FR-MKT-01", "Tutor service listings", [
            "A tutor shall create, edit, deactivate, and delete owned service listings with title, subject, level, price, mode, availability, description, and optional trial/media data.",
            "Only active services shall appear in public discovery unless an authorized administrator is reviewing records."
        ]),
        ("FR-MKT-02", "Student tutoring requests", [
            "A student shall create, update, close, and delete an owned tutoring request with subject, level, budget, location, mode, preferred time, experience preference, and description.",
            "Tutors shall browse open requests; public users may view published request summaries."
        ]),
        ("FR-MKT-03", "Applications and hiring", [
            "A tutor shall submit at most one proposal per student request, including message, expected fee, and available time.",
            "The owning student or an administrator shall accept or reject an application.",
            "Accepting an application shall transactionally reject competing pending proposals, mark the request as hired, and create one pending booking for that request."
        ]),
        ("FR-BOOK-01", "Booking and availability", [
            "A student shall view up to 31 days of tutor availability and create trial, one-time, weekly, or monthly class bookings with a future date, time, duration, and supported mode.",
            "The system shall prevent overlapping active bookings for either participant and shall allow no more than one non-cancelled trial per student/tutor pair.",
            "Authorized participants shall follow role-specific booking transitions; only tutors shall set meeting links or class locations."
        ]),
        ("FR-MSG-01", "Messaging and notifications", [
            "Authenticated users shall exchange persistent database-backed messages with active users, view conversation summaries, and retrieve recent message history and unread state.",
            "A recipient shall mark messages or generated notifications as read, individually or in bulk.",
            "Recipients shall flag inappropriate messages, and the system shall deliver in-app notifications for significant marketplace and learning events."
        ]),
        ("FR-LEARN-01", "Study materials", [
            "A tutor shall publish material metadata with title, subject, file URL, and an intended student or valid teaching relationship.",
            "Authorized students shall view relevant materials; tutors and administrators shall remove materials where permitted."
        ]),
        ("FR-LEARN-02", "Assignments", [
            "A tutor shall create an assignment for a student with title, description, and deadline.",
            "The assigned student shall submit text and/or a file URL.",
            "The tutor or administrator shall grade a submitted assignment with marks and feedback."
        ]),
        ("FR-LEARN-03", "Quizzes and progress", [
            "A tutor shall create quizzes with structured questions and total score.",
            "An eligible student shall attempt or reattempt a quiz; the system shall automatically calculate the percentage score and keep one current attempt record per student/quiz pair.",
            "The student dashboard shall derive progress from bookings, assignments, quizzes, and completed learning activity."
        ]),
        ("FR-FIN-01", "Simulated payments", [
            "A student shall create one mock payment for an owned confirmed or completed booking using a supported demonstration method.",
            "The server shall derive the payable amount and, when marked paid, record the amount, configured commission, tutor earning, status, and paid time.",
            "The system shall not claim that simulated payment records represent real financial settlement."
        ]),
        ("FR-FIN-02", "Earnings and withdrawals", [
            "A tutor shall view gross paid revenue, platform commission, net earnings, and relevant payment history.",
            "A tutor shall request withdrawal up to the available simulated balance.",
            "An administrator shall approve or reject withdrawal requests."
        ]),
        ("FR-TRUST-01", "Reviews and ratings", [
            "A booking participant shall submit one 1–5 star review per booking with an optional comment.",
            "The system shall update or derive a tutor's average rating and display it in discovery and profile views.",
            "Authorized moderation shall remove an invalid review."
        ]),
        ("FR-TRUST-02", "Tutor verification", [
            "A tutor shall submit certificate, institution, proof, and demo-video references for review.",
            "An administrator shall list pending verifications and record a verified or rejected decision with feedback.",
            "Only verified tutors shall receive the verified badge."
        ]),
        ("FR-SAFE-01", "Reports and moderation", [
            "An authenticated user shall submit a categorized report with a description and optional reported-user or message context.",
            "Administrators shall review, resolve, or dismiss reports and preserve a moderation record.",
            "Administrators shall activate or suspend user accounts and suspended users shall be denied authenticated access."
        ]),
        ("FR-ADMIN-01", "Administrative dashboard", [
            "The admin dashboard shall display aggregate counts and marketplace/financial summaries.",
            "Administrators shall list users, students, tutors, posts, requests, applications, bookings, payments, reviews, verifications, reports, and withdrawals as authorized.",
            "Administrative endpoints and screens shall be inaccessible to non-admin users."
        ]),
    ]
    groups = {
        "3.1.1 Authentication and Profiles": requirements[:3],
        "3.1.2 Discovery and Marketplace": requirements[3:7],
        "3.1.3 Applications, Bookings, and Communication": requirements[7:10],
        "3.1.4 Learning Management": requirements[10:13],
        "3.1.5 Finance": requirements[13:15],
        "3.1.6 Trust, Safety, and Administration": requirements[15:],
    }
    for group, reqs in groups.items():
        doc.add_heading(group, level=3)
        for req in reqs:
            add_req(doc, *req)

    doc.add_heading("3.1.7 Business Rules", level=3)
    add_table(doc, ["ID", "Rule"], [
        ("BR-01", "A user email shall be unique and every account shall have exactly one platform role."),
        ("BR-02", "Only tutors may publish services or proposals; only students may publish tutoring requests or initiate bookings."),
        ("BR-03", "One tutor may submit no more than one application to the same student request."),
        ("BR-04", "A booking created from a student request shall be unique for that request."),
        ("BR-05", "A payment shall be unique to a booking and shall calculate commission and tutor earning consistently."),
        ("BR-06", "A reviewer may submit no more than one review per booking; ratings must be integers from 1 to 5."),
        ("BR-07", "Tutor and student access to learning resources shall require a legitimate teaching relationship where applicable."),
        ("BR-08", "Only an administrator may decide verification, resolve reports, suspend users, or process withdrawals."),
    ], [0.8, 6.0])

    doc.add_heading("3.2 Non-Functional Requirements", level=2)
    nfrs = [
        ("NFR-PERF-01", "Performance", "Under normal academic demonstration load, 95% of ordinary read/write API requests should complete within 2 seconds, excluding network latency and media transfer."),
        ("NFR-PERF-02", "Performance", "Search input shall be debounced, and catalog endpoints shall use bounded pagination to avoid unbounded responses."),
        ("NFR-SEC-01", "Security", "Passwords shall be stored only as bcrypt-compatible hashes; protected endpoints shall validate JWT identity and role on every request."),
        ("NFR-SEC-02", "Security", "The server shall use Helmet, configured CORS, authentication rate limiting, request validation, and safe error responses."),
        ("NFR-SEC-03", "Security", "Secrets, database credentials, and JWT keys shall be supplied through environment variables and shall not be committed."),
        ("NFR-SEC-04", "Security", "Production traffic shall use HTTPS, and authorization shall be enforced server-side even when the UI hides an action."),
        ("NFR-REL-01", "Reliability", "Database constraints and application checks shall preserve referential integrity, uniqueness, and valid state transitions."),
        ("NFR-REL-02", "Reliability", "Unexpected errors shall be handled centrally, logged without exposing secrets, and returned using a consistent JSON error shape."),
        ("NFR-USAB-01", "Usability", "The interface shall remain usable from 375px mobile width through common desktop widths without horizontal page scrolling."),
        ("NFR-USAB-02", "Usability", "Forms shall provide human-readable validation feedback, and async screens shall provide loading, success, empty, and failure states."),
        ("NFR-ACC-01", "Accessibility", "Interactive controls shall be keyboard operable, maintain visible focus, use semantic labels, and provide text alternatives for meaningful images."),
        ("NFR-MAINT-01", "Maintainability", "Backend code shall preserve MVC separation and shared middleware/utilities; frontend code shall use reusable components, hooks, and centralized API configuration."),
        ("NFR-TEST-01", "Testability", "Critical authentication, booking, availability, payment, tutor search, and model behavior shall be verifiable through automated tests and the Postman collection."),
        ("NFR-PORT-01", "Portability", "The application shall run with documented npm commands on a Node-compatible environment with MySQL 8."),
        ("NFR-SCALE-01", "Scalability", "The API shall remain stateless between requests, use database connection pooling, and bound collection responses so that additional application instances can be introduced without redesigning client contracts."),
        ("NFR-PRIV-01", "Privacy", "Profile and report data shall be disclosed only to users and roles with a legitimate platform need; public APIs shall avoid password hashes and moderation-only fields."),
    ]
    nfr_groups = [
        ("3.2.1 Performance Requirements", {"Performance"}),
        ("3.2.2 Security and Privacy Requirements", {"Security", "Privacy"}),
        ("3.2.3 Reliability and Availability Requirements", {"Reliability"}),
        ("3.2.4 Usability and Accessibility Requirements", {"Usability", "Accessibility"}),
        (
            "3.2.5 Maintainability, Testability, Portability, and Scalability",
            {"Maintainability", "Testability", "Portability", "Scalability"},
        ),
    ]
    for heading, categories in nfr_groups:
        heading_paragraph = doc.add_heading(heading, level=3)
        if heading.startswith("3.2.3"):
            heading_paragraph.paragraph_format.page_break_before = True
        rows = [(req_id, requirement) for req_id, category, requirement in nfrs if category in categories]
        add_table(doc, ["ID", "Requirement"], rows, [1.25, 5.55])

    doc.add_heading("3.3 External Interface Requirements", level=2)
    doc.add_heading("3.3.1 User Interfaces", level=3)
    for item in [
        "Public pages shall provide consistent navigation, marketplace search entry points, responsive tutor/service cards, and clear login and registration actions.",
        "Authenticated workspaces shall provide role-specific dashboards and navigation. Unauthorized routes shall redirect or deny access.",
        "Forms shall show labels, required states, inline validation, loading indicators, success/error feedback, and accessible focus behavior.",
        "Catalog pages shall support pagination or progressive loading and clear empty states.",
        "The mobile student experience shall include bottom navigation and a searchable quick-jump interface.",
        "Dialogs, controls, and media shall support keyboard operation, meaningful accessible names, and reduced-motion preferences.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("3.3.2 Hardware Interfaces", level=3)
    doc.add_paragraph(
        "No specialized hardware is required. The system shall operate on ordinary client devices capable of running a modern "
        "web browser and on standard Linux-compatible or cloud server infrastructure."
    )
    doc.add_heading("3.3.3 Software Interfaces", level=3)
    add_table(doc, ["Interface", "Purpose", "Data format"], [
        ("React ↔ Express API", "Authentication, discovery, marketplace, learning, finance, and administration operations.", "HTTPS + JSON"),
        ("Express ↔ MySQL", "Transactional persistence, filtering, aggregation, referential integrity, and search.", "SQL via mysql2"),
        ("Browser media", "Bundled demo videos, thumbnails, avatars, and remote URL references.", "Web media / URL"),
        ("Postman collection", "Manual API exploration and end-to-end workflow verification.", "HTTP + JSON"),
    ], [1.55, 3.75, 1.5])
    doc.add_heading("3.3.4 Communication Interfaces", level=3)
    doc.add_paragraph(
        "The client shall communicate with the REST API over HTTP/HTTPS and JSON. Authenticated requests shall include a bearer "
        "JWT. The API shall return consistent success and error envelopes with appropriate HTTP status codes and request identifiers. "
        "CORS shall permit only configured origins, and production deployments shall use HTTPS."
    )

    doc.add_heading("4. Technology Stack & Architectural Overview", level=1)
    doc.add_heading("4.1 Technology Stack Components", level=2)
    add_table(doc, ["Layer / concern", "Technology", "Purpose"], [
        ("Frontend", "React 19, React Router 7, Axios, Vite 7", "Single-page UI, routing, HTTP access, and production bundling"),
        ("Backend", "Node.js, Express 5, ES modules", "REST API, business workflows, validation, and authorization"),
        ("Database", "MySQL 8 / MariaDB 12+, mysql2", "Normalized relational persistence and connection pooling"),
        ("Security", "JWT, bcryptjs, Helmet, CORS, rate limiting", "Session identity, password hashing, headers, origin control, and abuse resistance"),
        ("Quality", "Node test runner, Supertest, Postman", "Automated and manual verification"),
    ], [1.25, 2.4, 3.15])
    doc.add_heading("4.2 High-Level Architecture", level=2)
    add_table(doc, ["Layer", "Components", "Responsibility"], [
        ("Presentation", "React pages, layouts, components, hooks, context, Axios", "Responsive UI, navigation, form state, protected routes, and feedback"),
        ("Application/API", "Express routes, middleware, controllers, utilities", "Authentication, authorization, validation, workflow, response handling"),
        ("Domain/Data", "Models, query builders, MySQL pool", "Persistence, search, aggregation, constraints, and transactions"),
        ("Quality/Operations", "Node tests, Supertest, Postman, npm scripts", "Verification, reproducible setup, database lifecycle, and builds"),
    ], [1.25, 2.7, 2.85])
    doc.add_heading("4.3 Data Model and Integrity", level=2)
    doc.add_heading("4.3.1 Core Entities", level=3)
    add_table(doc, ["Entity", "Purpose", "Principal relationships"], [
        ("User", "Identity, role, contact, account state", "Owns one student/tutor profile; participates throughout the system"),
        ("StudentProfile / TutorProfile", "Role-specific profile and discovery attributes", "One-to-one with User"),
        ("TutorPost", "Marketed tutoring service/course", "Tutor → many posts; referenced by bookings and engagement"),
        ("StudentRequest", "Published learning demand", "Student → many requests; receives applications"),
        ("Application", "Tutor proposal to a request", "Joins Tutor and StudentRequest"),
        ("Booking", "Scheduled learning relationship", "Joins student, tutor, optional post/request"),
        ("Message / Notification", "Communication and event awareness", "Associated with users and optionally bookings"),
        ("Material / Assignment / Quiz", "Learning content and assessment", "Tutor-to-student learning relationship"),
        ("Payment / Withdrawal", "Simulated financial ledger", "Payment belongs to booking; withdrawal belongs to tutor"),
        ("Review / Verification / Report", "Trust, quality, and moderation", "Associated with users, bookings, and administrators"),
        ("SavedTutor / SavedCourse / CourseView", "Discovery engagement", "Student-to-tutor/course interactions"),
    ], [1.6, 2.4, 2.8])
    doc.add_heading("4.3.2 Relationship Model", level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "USER ──1:1── PROFILE\n"
        "STUDENT ──1:N── REQUEST ──1:N── APPLICATION ──N:1── TUTOR\n"
        "STUDENT ──1:N── BOOKING ──N:1── TUTOR\n"
        "BOOKING ──1:0..1── PAYMENT   |   BOOKING ──1:N── REVIEW\n"
        "TUTOR ──1:N── POST   |   TUTOR ↔ STUDENT: MESSAGE, MATERIAL, ASSIGNMENT, QUIZ"
    )
    apply_font(r, name=MONO_FONT, size=9.5)
    shade_p = OxmlElement("w:shd")
    shade_p.set(qn("w:fill"), PALE)
    p._p.get_or_add_pPr().append(shade_p)
    doc.add_heading("4.3.3 Data Integrity and Retention", level=3)
    for item in [
        "Foreign keys shall apply cascade, restrict, or set-null behavior appropriate to the historical value of each record.",
        "Unique indexes shall protect email, duplicate applications, request-derived bookings, booking payments, reviews, quiz attempts, and saved relationships.",
        "Status values shall use controlled enumerations. Monetary values shall use fixed-precision decimal types.",
        "Operational records should be retained for the duration of the academic demonstration; a production deployment would require a formal retention and deletion policy.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4.4 Key Use Cases", level=2)
    use_cases = [
        ("UC-01 Find and book a tutor", "Student", "Student is authenticated; tutor/service is active.", "Search/filter → open details → choose class type/time → check availability → create booking.", "Pending booking is stored without conflict."),
        ("UC-02 Hire from a request", "Student, Tutor", "Student request is open.", "Tutor submits proposal → student reviews → accepts proposal → competing proposals are rejected → pending booking is created.", "Application is accepted and request is hired."),
        ("UC-03 Complete learning work", "Tutor, Student", "Valid learning relationship exists.", "Tutor posts material/assignment/quiz → student views/submits/attempts → tutor grades assignment.", "Scores and progress become available."),
        ("UC-04 Pay and view earnings", "Student, Tutor", "Booking exists and is payable.", "Student creates mock payment → marks it paid → commission is calculated → tutor views earnings.", "Ledger shows paid amount, commission, and net earning."),
        ("UC-05 Verify a tutor", "Tutor, Admin", "Tutor is authenticated.", "Tutor submits evidence URLs → admin reviews → verifies or rejects with feedback.", "Verification status and badge eligibility are updated."),
        ("UC-06 Moderate a report", "User, Admin", "User is authenticated.", "User submits report → admin filters/reviews → resolves or dismisses → optionally suspends user.", "Moderation status and action are recorded."),
    ]
    add_table(doc, ["Use case", "Actors", "Precondition", "Main flow", "Postcondition"], use_cases, [1.25, 0.8, 1.3, 2.25, 1.2])
    doc.add_heading("4.4.1 Booking State Model", level=3)
    doc.add_paragraph(
        "A booking begins as pending. A tutor may confirm it and, after the scheduled class, mark it completed. Students and "
        "tutors may cancel or reschedule active bookings within their permitted transitions; administrators may cancel active "
        "bookings. Meeting details are tutor-controlled. The system shall reject terminal-state changes, premature completion, "
        "role violations, and rescheduling that conflicts with either participant's active bookings."
    )

    doc.add_heading("5. Tentative Development Plan (Agile Methodology)", level=1)
    doc.add_paragraph(
        "The following eight-sprint plan mirrors the academic planning style of the reference document while mapping directly "
        "to Mentor Market's modules. Each sprint should end with demonstrable, tested increments."
    )
    add_table(doc, ["Sprint", "Focus", "Principal deliverables"], [
        ("1", "Foundation", "Repository structure, environment configuration, MySQL schema/seed, Express and React shells"),
        ("2", "Identity and profiles", "Registration/login, JWT/RBAC, student and tutor profiles, protected layouts"),
        ("3", "Marketplace discovery", "Tutor services, student requests, search/filter, public details, saved items"),
        ("4", "Applications and bookings", "Proposals, decisions, availability, booking workflow, schedule conflict protection"),
        ("5", "Communication and learning", "Messages, notifications, materials, assignments, submissions, grading"),
        ("6", "Assessment and progress", "Quizzes, automatic scoring, course engagement, progress analytics"),
        ("7", "Finance, trust, and admin", "Mock payments, commission, earnings, withdrawals, reviews, verification, reports, admin dashboard"),
        ("8", "Hardening and delivery", "Responsive/accessibility pass, security review, automated tests, Postman, documentation, final demo"),
    ], [0.7, 1.65, 4.45])
    doc.add_heading("5.1 Definition of Done", level=2)
    for item in [
        "Acceptance behavior is implemented and demonstrated for permitted roles.",
        "Validation, authorization, loading, empty, and error states are handled.",
        "Database changes include safe, repeatable schema or migration updates.",
        "Relevant automated tests pass and regressions are not introduced.",
        "User-facing behavior and API usage are documented where necessary.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("6. Acceptance Criteria", level=1)
    doc.add_heading("6.1 System Acceptance Criteria", level=2)
    criteria = [
        "The database schema and seed data initialize successfully on MySQL 8 or the documented compatible local MariaDB environment.",
        "Student, tutor, and administrator users can authenticate and are restricted to their authorized routes and operations.",
        "A student can discover a tutor, create a booking without schedule conflict, use the mock payment ledger, and submit a valid completed-booking review.",
        "A tutor can publish a service, apply to a student request, manage teaching resources and assessments, and view calculated earnings.",
        "An administrator can inspect analytics, moderate users and reports, decide tutor verification, and process withdrawal requests through protected backend operations.",
        "Backend automated tests pass and the frontend production build completes without errors.",
        "The main public and authenticated workflows remain usable on mobile and desktop layouts.",
        "Finance and file-reference features are clearly identified as simulations or URL-based MVP behavior and do not imply real settlement or managed upload.",
    ]
    for criterion in criteria:
        add_number(doc, criterion)

    doc.add_heading("6.2 Requirement Traceability Matrix", level=2)
    add_table(doc, ["Requirement group", "Primary implementation area", "Acceptance evidence"], [
        ("FR-AUTH / FR-PROF", "authRoutes, userRoutes, student/tutor profiles, AuthContext", "Registration/login tests; protected route checks; profile update"),
        ("FR-DISC", "tutor search query builder, discovery/course pages, engagement routes", "Filter/search results; saved/recent data; frontend build"),
        ("FR-MKT", "post, request, and application routes/controllers", "CRUD role checks; application uniqueness and transactional status"),
        ("FR-BOOK", "booking controller, availability utilities, bookings table", "Availability/conflict and transition tests"),
        ("FR-MSG", "message and notification modules", "Persistent conversations, unread state, message flagging, and in-app notifications"),
        ("FR-LEARN", "materials, assignment, quiz, and progress modules", "Create/submit/grade/attempt flows and score output"),
        ("FR-FIN", "payment, earnings, withdrawal modules", "Payment route tests; server-side amount/commission calculation; admin decision"),
        ("FR-TRUST / FR-SAFE", "reviews, verification, report, admin modules", "Rating constraints; decision permissions; suspension behavior"),
        ("NFR-SEC / NFR-TEST", "middleware, tests, Postman, npm scripts", "Automated test pass; production build; HTTP/security review"),
    ], [1.35, 3.0, 2.45])

    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "This SRS establishes a project-specific baseline for Mentor Market: a responsive three-role tutoring marketplace that "
        "supports discovery, matching, bookings, communication, learning activities, simulated finance, trust, and administration. "
        "Its requirements reflect the implemented React, Express, and MySQL system while clearly separating MVP simulations and "
        "future integrations. Following the stated requirements, acceptance criteria, and sprint plan will support consistent "
        "testing, evaluation, maintenance, and future development."
    )

    doc.add_heading("Appendix A — Out of Scope for the MVP", level=1)
    for item in [
        "Real payment gateway processing, bank settlement, chargebacks, and identity/KYC checks.",
        "Managed binary file upload, transcoding, antivirus scanning, or content delivery storage.",
        "Native iOS/Android applications, offline-first synchronization, or SMS/push delivery.",
        "Live video classroom hosting, calendar provider synchronization, or automated attendance.",
        "Automated credential verification against external institutions.",
        "Production-scale recommendation machine learning; current recommendations are rule/data driven.",
    ]:
        add_bullet(doc, item)
    appendix_b = doc.add_heading("Appendix B — Risks and Mitigations", level=1)
    appendix_b.paragraph_format.page_break_before = True
    add_table(doc, ["Risk", "Impact", "Mitigation"], [
        ("Schedule race conditions", "Double-booked tutor", "Use availability checks, database constraints/transactions, and concurrency tests."),
        ("Broken authorization", "Cross-role data exposure or mutation", "Enforce middleware and ownership checks on the server; test denied paths."),
        ("Untrusted URLs/content", "Unsafe links or inappropriate media", "Validate URL shape, restrict schemes, sanitize display, and provide reporting/moderation."),
        ("Mock payment confusion", "Users mistake demo ledger for real settlement", "Label all finance screens and records as simulated."),
        ("Large catalog queries", "Slow discovery and high database load", "Use indexes, bounded pagination, debouncing, and query profiling."),
        ("Incomplete tutor evidence", "Weak marketplace trust", "Expose verification status clearly and require manual admin review."),
    ], [1.5, 1.9, 3.4])

    props = doc.core_properties
    props.title = "Mentor Market Software Requirements Specification"
    props.subject = "Project-specific SRS for the Mentor Market EdTech tutoring marketplace"
    props.author = "Mentor Market Project Team"
    props.last_modified_by = "Mentor Market Project Team"
    props.keywords = "Mentor Market, SRS, software requirements, EdTech, tutoring marketplace"
    props.comments = "Generated from the implemented Mentor Market repository and the supplied academic SRS structure."
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
