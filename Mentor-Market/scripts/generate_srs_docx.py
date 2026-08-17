"""Generate the reference-styled Mentor Market CSE470 SRS.

The supplied WeHeal PDF is intentionally treated as a visual template:
US Letter paper, one-inch margins, Times New Roman body text, black formal
pages, blue hyperlinks, a centered cover table, an untitled TOC, and a
screenshot-based sprint appendix.
"""

from pathlib import Path
import sys


sys.path.insert(0, "/tmp/mentor_market_docx")

from PIL import Image
from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "documentation"
OUTPUT = OUTPUT_DIR / "Mentor_Market_SRS.docx"
ASSET_DIR = OUTPUT_DIR / "srs_assets"

TIMES = "Times New Roman"
ARIAL = "Arial"
BLACK = "000000"
LINK_BLUE = "1154CC"
FEATURE_BLUE = "0000FF"
RED = "FF0000"
DEEP_RED = "CC0000"
WHITE = "FFFFFF"


TEAM = [
    ("23201560", "Sieuti Zaman"),
    ("23201291", "Mumtasim Daiyan"),
    ("23201199", "Abir Saha"),
    ("23201451", "Samina Fairooz Urbi"),
]


def set_run_font(
    run,
    name=TIMES,
    size=11,
    bold=False,
    italic=False,
    color=BLACK,
    underline=False,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_style_font(style, name, size, bold=False, color=BLACK):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_repeatable_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def configure_document(doc):
    for section in doc.sections:
        set_repeatable_page_geometry(section)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, TIMES, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)

    h1 = styles["Heading 1"]
    set_style_font(h1, TIMES, 17, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.line_spacing = 1
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(17)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, TIMES, 13, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, TIMES, 11, bold=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 1
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    # Preserve the reference's absence of visible headers, footers, and page numbers.
    for section in doc.sections:
        section.header.paragraphs[0].clear()
        section.footer.paragraphs[0].clear()

    compat = doc.settings._element.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        doc.settings._element.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)


def add_h1(doc, text, page_break=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = page_break
    set_run_font(p.add_run(text), size=17, bold=True)
    return p


def add_h2(doc, text, page_break=False):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.page_break_before = page_break
    set_run_font(p.add_run(text), size=13, bold=True)
    return p


def add_h3(doc, text, page_break=False):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.page_break_before = page_break
    set_run_font(p.add_run(text), size=11, bold=True)
    return p


def add_body(doc, text, bold_prefix=None, after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix) :]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_rich_body(doc, parts, after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(after)
    for part in parts:
        if isinstance(part, str):
            set_run_font(p.add_run(part))
        else:
            text, bold, italic, color = part
            set_run_font(
                p.add_run(text),
                bold=bold,
                italic=italic,
                color=color,
            )
    return p


def add_manual_list(
    doc,
    marker,
    text,
    label=None,
    level=1,
    after=0,
    marker_font=ARIAL,
    size=11,
    color=BLACK,
):
    p = doc.add_paragraph()
    text_indent = 0.5 + (level - 1) * 0.5
    p.paragraph_format.left_indent = Inches(text_indent)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(text_indent), WD_TAB_ALIGNMENT.LEFT)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run(marker), name=marker_font, size=size, color=color)
    p.add_run("\t")
    if label:
        set_run_font(p.add_run(label), size=size, bold=True, color=color)
    set_run_font(p.add_run(text), size=size, color=color)
    return p


def add_bullet(doc, text, label=None, level=1, after=0):
    return add_manual_list(doc, "●", text, label=label, level=level, after=after)


def add_circle_bullet(doc, text, label=None, level=2, after=0):
    return add_manual_list(doc, "○", text, label=label, level=level, after=after)


def add_numbered(doc, number, text, label=None, after=0, level=1):
    return add_manual_list(doc, f"{number}.", text, label=label, level=level, after=after)


def add_external_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), TIMES)
    rfonts.set(qn("w:hAnsi"), TIMES)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    rpr.append(size)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([rpr, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_margins(cell, top=0, start=70, bottom=0, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, size=8, color=BLACK):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_cover(doc):
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(153.5)
    name_p.paragraph_format.space_after = Pt(9)
    name_p.paragraph_format.line_spacing = 1
    set_run_font(name_p.add_run("“Mentor Market”"), size=26, bold=True, italic=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(23)
    subtitle_p.paragraph_format.line_spacing = 1
    set_run_font(
        subtitle_p.add_run("an all-in-one tutoring web application"),
        size=26,
        bold=True,
        italic=True,
    )

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(21)
    title_p.paragraph_format.line_spacing = 1
    set_run_font(title_p.add_run("Software Requirements Specification"), size=14, bold=True)

    prepared_p = doc.add_paragraph()
    prepared_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_p.paragraph_format.space_after = Pt(38)
    prepared_p.paragraph_format.line_spacing = 1
    set_run_font(prepared_p.add_run("Prepared by"), size=14, bold=True)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "6340")
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    set_table_borders(table, size=8)

    rows = [("Student ID", "Name"), *TEAM]
    widths = (2810, 3530)
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(width))
    for row_index, (student_id, name) in enumerate(rows):
        row = table.rows[row_index]
        row.height = Pt(21)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell_index, value in enumerate((student_id, name)):
            cell = row.cells[cell_index]
            set_cell_width(cell, widths[cell_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            set_run_font(p.add_run(value), size=12, bold=row_index == 0)


TOC_ENTRIES = [
    ("1. Introduction", 1, 3),
    ("1.1 Purpose", 2, 3),
    ("1.2 Scope", 2, 3),
    ("1.3 Definitions, Acronyms, and Abbreviations", 2, 3),
    ("1.4 References", 2, 3),
    ("1.5 Overview", 2, 4),
    ("2. Overall Description", 1, 4),
    ("2.1 Product Perspective", 2, 4),
    ("2.2 Product Features", 2, 4),
    ("2.3 User Classes and Characteristics", 2, 4),
    ("2.4 Operating Environment", 2, 5),
    ("2.5 Constraints", 2, 5),
    ("2.6 Assumptions and Dependencies", 2, 5),
    ("3. System Requirements", 1, 5),
    ("3.1 Functional Requirements", 2, 5),
    ("3.1.1 Authentication & Authorization", 3, 5),
    ("3.1.2 Student & Tutor Interaction", 3, 6),
    ("3.1.3 Learning Management Services", 3, 6),
    ("3.1.4 Additional Services", 3, 7),
    ("3.2 Non-Functional Requirements", 2, 7),
    ("3.2.1 Performance Requirements", 3, 7),
    ("3.2.2 Security Requirements", 3, 7),
    ("3.2.3 Reliability & Availability", 3, 7),
    ("3.2.4 Maintainability", 3, 7),
    ("3.2.5 Scalability", 3, 7),
    ("3.3 External Interface Requirements", 2, 8),
    ("3.3.1 User Interfaces", 3, 8),
    ("3.3.2 Hardware Interfaces", 3, 8),
    ("3.3.3 Software Interfaces", 3, 8),
    ("3.3.4 Communication Interfaces", 3, 8),
    ("4. Technology Stack & Architectural Overview", 1, 8),
    ("4.1 Technology Stack Components", 2, 8),
    ("4.2 High-Level Architecture", 2, 8),
    ("5. Tentative Development Plan (Agile Methodology)", 1, 9),
    ("Sprint 1 (Weeks 1–2): Foundation, Authentication & Profiles", 2, 9),
    ("Sprint 2 (Weeks 3–4): Marketplace & Discovery", 2, 9),
    ("Sprint 3 (Weeks 5–6): Applications, Bookings & Communication", 2, 9),
    ("Sprint 4 (Week 7): Learning, Finance & Trust Services", 2, 9),
    ("Sprint 5 (Weeks 8–9): Administration, Quality & Delivery", 2, 9),
    ("6. Acceptance Criteria", 1, 10),
    ("7. Conclusion", 1, 10),
]


def add_toc(doc):
    for index, (title, level, page) in enumerate(TOC_ENTRIES):
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = index == 0
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        p.paragraph_format.right_indent = Inches(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(15.55)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        set_run_font(p.add_run(title), size=11, bold=level == 1)
        set_run_font(p.add_run(f"\t{page}"), size=11, bold=level == 1)


def add_reference_link(doc, text, url):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1), WD_TAB_ALIGNMENT.LEFT)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("○"), name=ARIAL)
    p.add_run("\t")
    add_external_hyperlink(p, text, url)
    return p


def add_formal_srs(doc):
    # Physical page 3
    add_h1(doc, "1. Introduction", page_break=True)
    add_h2(doc, "1.1 Purpose")
    add_rich_body(
        doc,
        [
            "This Software Requirements Specification (SRS) document outlines the requirements for developing ",
            ("“Mentor Market” - an all-in-one tutoring marketplace web application", True, False, BLACK),
            " using React.js, Node.js, Express.js, and MySQL. The system connects students with tutors and supports their learning relationship from discovery and hiring through classes, assessment, simulated payment, progress, and review.",
        ],
        after=9,
    )

    add_h2(doc, "1.2 Scope")
    add_body(
        doc,
        "The scope of this project includes the design, development, testing, and delivery of the Mentor Market web application, catering to:",
        after=6,
    )
    add_bullet(doc, " who can browse the public marketplace and register or log in.", label="Visitors")
    add_bullet(doc, " who can discover tutors, publish requests, book classes, learn, pay in the mock ledger, and review.", label="Students")
    add_bullet(doc, " who can publish services, apply to requests, teach, assess students, and monitor simulated earnings.", label="Tutors")
    add_bullet(doc, " who oversee users, verification, reports, records, and platform analytics.", label="Administrators", after=6)
    add_body(
        doc,
        "Mentor Market is an academic MVP. Payments and withdrawals are simulated, files are represented by URLs, messages are database-backed, and the platform does not host real-time video classes.",
        after=8,
    )

    add_h2(doc, "1.3 Definitions, Acronyms, and Abbreviations")
    definitions = [
        ("API", ": Application Programming Interface."),
        ("JWT", ": JSON Web Token used for authenticated sessions."),
        ("MVC", ": Model–View–Controller backend organization."),
        ("MVP", ": Minimum Viable Product."),
        ("RBAC", ": Role-Based Access Control."),
        ("REST", ": Representational State Transfer style used by the HTTP API."),
    ]
    for index, (label, text) in enumerate(definitions):
        add_bullet(doc, text, label=label, after=0 if index < len(definitions) - 1 else 7)

    add_h2(doc, "1.4 References")
    add_bullet(doc, " documentation:", label="Project and technology")
    add_reference_link(doc, "Mentor Market repository documentation", "https://github.com/")
    add_reference_link(doc, "React.js", "https://react.dev/")
    add_reference_link(doc, "Express.js", "https://expressjs.com/")
    add_reference_link(doc, "Node.js", "https://nodejs.org/")
    add_reference_link(doc, "MySQL", "https://dev.mysql.com/doc/")

    # Physical page 4
    add_h2(doc, "1.5 Overview", page_break=True)
    add_body(
        doc,
        "Section 2 provides an overall description of the product, including its user classes and environment. Section 3 details the functional and non-functional requirements. Section 4 outlines the technology stack and architectural overview. Section 5 presents the development plan, including sprint-wise breakdown and deliverables.",
        after=14,
    )

    add_h1(doc, "2. Overall Description")
    add_h2(doc, "2.1 Product Perspective")
    add_body(
        doc,
        "Mentor Market is a standalone three-tier web application. A React single-page interface communicates with an Express.js REST API using HTTP and JSON. The API applies authentication, role checks, validation, ownership rules, and business workflows before accessing the MySQL database.",
        after=9,
    )

    add_h2(doc, "2.2 Product Features")
    features = [
        ("Identity and Role Security", ": Registration, login, JWT sessions, profiles, protected routes, and role dashboards."),
        ("Tutor and Course Discovery", ": Search, filters, details, ratings, verification badges, saved items, history, and comparison."),
        ("Student Tutoring Requests", ": Students publish and manage learning needs with budget, mode, location, and schedule."),
        ("Tutor Service Publishing", ": Tutors manage priced services with subject, level, mode, availability, trial, and media URLs."),
        ("Applications and Matching", ": Tutors propose; students decide; acceptance hires one tutor and creates a booking."),
        ("Scheduling and Booking", ": Availability, trial and recurring classes, conflict prevention, and controlled status transitions."),
        ("Communication and Notifications", ": Persistent messages, unread state, reporting, and database-backed in-app notices."),
        ("Learning Management and Progress", ": Materials, assignments, grading, quizzes, automatic scores, and analytics."),
        ("Simulated Finance", ": Mock payments, platform commission, tutor earnings, and withdrawal requests."),
        ("Trust, Safety, and Administration", ": Reviews, tutor verification, reports, suspension, moderation, and analytics."),
    ]
    for number, (label, text) in enumerate(features, 1):
        add_numbered(doc, number, text, label=label, after=0)

    add_h2(doc, "2.3 User Classes and Characteristics")
    add_numbered(
        doc,
        1,
        ": Unauthenticated users who browse public tutors, services, student requests, and information pages.",
        label="Visitors",
        after=0,
    )

    # Physical page 5
    p = add_numbered(
        doc,
        2,
        ": Learners seeking discovery, booking, communication, learning, simulated payment, and progress tools.",
        label="Students",
        after=0,
    )
    p.paragraph_format.page_break_before = True
    add_numbered(
        doc,
        3,
        ": Mentors who publish services, respond to requests, manage classes, assess learners, and view earnings.",
        label="Tutors",
        after=0,
    )
    add_numbered(
        doc,
        4,
        ": Platform overseers responsible for analytics, trust, user status, and marketplace moderation.",
        label="Administrators",
        after=7,
    )

    add_h2(doc, "2.4 Operating Environment")
    environment = [
        ("Web Application", ": Accessible on modern Chrome, Firefox, Edge, and Safari browsers."),
        ("Mobile Compatibility", ": Responsive layouts for smartphones and tablets."),
        ("Server-Side", ": Node.js and Express.js on a Linux-compatible host."),
        ("Database", ": MySQL 8 through mysql2 connection pooling."),
        ("Development", ": npm, Visual Studio Code, Node test runner, Supertest, and Postman."),
    ]
    for label, text in environment:
        add_bullet(doc, text, label=label)

    add_h2(doc, "2.5 Constraints")
    constraints = [
        ("Security", ": Passwords use bcrypt-compatible hashes; protected actions require JWT and server RBAC."),
        ("Relational Data", ": Persistence must remain compatible with the normalized MySQL schema."),
        ("MVP Finance", ": Payments, commission, earnings, and withdrawals are simulations only."),
        ("MVP Files and Media", ": Files and videos are URL references; live video hosting is outside scope."),
        ("Time & Resources", ": Delivery is limited to the CSE470 academic schedule and team resources."),
    ]
    for label, text in constraints:
        add_bullet(doc, text, label=label)

    add_h2(doc, "2.6 Assumptions and Dependencies")
    assumptions = [
        "Users have a stable internet connection and supported web browser.",
        "Tutors provide accurate profile, schedule, price, and credential information.",
        "Administrators manually review verification evidence, safety reports, and withdrawals.",
        "The Node.js runtime, MySQL service, JWT secret, API URL, and CORS origins are configured.",
    ]
    for text in assumptions:
        add_bullet(doc, text)

    add_h1(doc, "3. System Requirements")
    add_h2(doc, "3.1 Functional Requirements")
    add_h3(doc, "3.1.1 Authentication & Authorization")
    add_numbered(doc, 1, " Account Authentication:", after=0)
    add_circle_bullet(
        doc,
        "FR-1: The system shall allow a visitor to register as a student or tutor using a valid unique email and password.",
    )
    add_circle_bullet(
        doc,
        "FR-2: The system shall hash passwords and authenticate active users before issuing a time-limited JWT.",
    )

    # Physical page 6
    p = add_numbered(doc, 2, " Role-Based Access Control (RBAC):", after=0)
    p.paragraph_format.page_break_before = True
    add_circle_bullet(
        doc,
        "FR-3: The system shall provide separate student, tutor, and administrator dashboards.",
    )
    add_circle_bullet(
        doc,
        "FR-4: The system shall enforce role, ownership, and active-account checks on every protected operation.",
        after=5,
    )

    add_h3(doc, "3.1.2 Student & Tutor Interaction")
    interaction_groups = [
        (
            "Tutor Discovery:",
            [
                "FR-5: Visitors and students shall search and filter active tutors by text, subject, mode, location, price, rating, and day.",
                "FR-6: Students shall save tutors and courses, record recent course views, and compare selected courses.",
            ],
        ),
        (
            "Marketplace Publishing:",
            [
                "FR-7: Tutors shall create, update, deactivate, and delete their own service listings.",
                "FR-8: Students shall create, update, close, and delete their own tutoring requests.",
            ],
        ),
        (
            "Applications & Bookings:",
            [
                "FR-9: A tutor shall submit one proposal per open request; acceptance shall reject competitors and create one pending booking.",
                "FR-10: Students shall request supported classes only within tutor availability and without active schedule conflicts.",
                "FR-11: Participants shall follow role-permitted booking transitions; only tutors shall update class access details after creation.",
            ],
        ),
    ]
    for group_index, (label, requirements) in enumerate(interaction_groups, 1):
        add_numbered(doc, group_index, f" {label}")
        for requirement in requirements:
            add_circle_bullet(doc, requirement)

    add_h3(doc, "3.1.3 Learning Management Services")
    learning_groups = [
        (
            "Communication:",
            [
                "FR-12: Authenticated users shall exchange persistent messages, view conversations, unread counts, and recent history.",
                "FR-13: The system shall create database-backed notifications for significant marketplace and learning events.",
            ],
        ),
        (
            "Materials & Assignments:",
            [
                "FR-14: Tutors shall share URL-based study materials with eligible students.",
                "FR-15: Tutors shall create assignments for students with a valid teaching relationship.",
                "FR-16: Students shall submit assignment text or a file URL, and tutors shall grade submitted work.",
            ],
        ),
        (
            "Quizzes & Progress:",
            [
                "FR-17: Tutors shall create structured quizzes; eligible students shall attempt them and receive automatic percentage scores.",
                "FR-18: The student dashboard shall derive progress from classes, assignments, grades, quizzes, and tutor feedback.",
            ],
        ),
    ]
    for group_index, (label, requirements) in enumerate(learning_groups, 1):
        add_numbered(doc, group_index, f" {label}")
        for requirement in requirements:
            add_circle_bullet(doc, requirement)

    # Physical page 7
    add_h3(doc, "3.1.4 Additional Services", page_break=True)
    additional_groups = [
        (
            "Reviews & Simulated Finance:",
            [
                "FR-19: A booking participant shall submit one 1–5 rating for a completed booking.",
                "FR-20: The server shall create one mock booking payment and derive its amount, commission, and tutor earning.",
                "FR-21: Tutors shall view simulated earnings and request withdrawals within their available balance.",
            ],
        ),
        (
            "Trust & Safety:",
            [
                "FR-22: Tutors shall submit credential URLs; administrators shall verify or reject them with feedback.",
                "FR-23: Users shall submit safety reports; administrators shall investigate, resolve, or dismiss them.",
            ],
        ),
        (
            "Administration:",
            [
                "FR-24: Administrators shall view analytics and protected marketplace records.",
                "FR-25: Administrators shall activate or suspend users but shall not suspend their own account.",
            ],
        ),
    ]
    for group_index, (label, requirements) in enumerate(additional_groups, 1):
        add_numbered(doc, group_index, f" {label}")
        for requirement in requirements:
            add_circle_bullet(doc, requirement)

    add_h2(doc, "3.2 Non-Functional Requirements")
    nfr_groups = [
        (
            "3.2.1 Performance Requirements",
            [
                "NFR-1: Normal API requests should complete within two seconds, excluding network and remote-media delay.",
                "NFR-2: Search shall use debouncing where appropriate, and collection responses shall remain bounded.",
            ],
        ),
        (
            "3.2.2 Security Requirements",
            [
                "NFR-3: Passwords shall be hashed; protected requests shall validate JWT, role, status, and ownership.",
                "NFR-4: The backend shall use validation, secure headers, CORS rules, rate limiting, parameterized SQL, and controlled responses.",
                "NFR-5: Production shall use HTTPS and environment-supplied secrets.",
            ],
        ),
        (
            "3.2.3 Reliability & Availability",
            [
                "NFR-6: Constraints and transactions shall protect integrity, uniqueness, balances, and valid state changes.",
                "NFR-7: The API shall expose health and consistent errors; deployment owns backups and uptime.",
            ],
        ),
        (
            "3.2.4 Maintainability",
            [
                "NFR-8: Backend MVC layers and reusable frontend components, hooks, and API configuration shall stay separated.",
                "NFR-9: Setup, schema, API examples, and critical tests shall remain current.",
            ],
        ),
        (
            "3.2.5 Scalability",
            [
                "NFR-10: The stateless API shall use MySQL pooling, indexes, reusable query builders, and bounded results so instances can scale without changing client contracts.",
            ],
        ),
    ]
    for heading, requirements in nfr_groups:
        add_h3(doc, heading)
        for requirement in requirements:
            add_bullet(doc, requirement)

    # Physical page 8
    add_h2(doc, "3.3 External Interface Requirements", page_break=True)
    add_h3(doc, "3.3.1 User Interfaces")
    add_bullet(doc, "UI-1: A responsive React interface shall support desktop, tablet, and mobile viewports.")
    add_bullet(doc, "UI-2: Public navigation and separate role dashboards shall show appropriate actions and feedback.")
    add_bullet(doc, "UI-3: Forms and dialogs shall support labels, keyboard focus, validation, loading, empty, and error states.")

    add_h3(doc, "3.3.2 Hardware Interfaces")
    add_bullet(doc, "HI-1: No specialized hardware is required beyond a modern browser client and standard server infrastructure.")

    add_h3(doc, "3.3.3 Software Interfaces")
    add_bullet(doc, "SI-1: Express.js shall communicate with MySQL through mysql2 for queries and transactions.")
    add_bullet(doc, "SI-2: The browser shall display bundled media and remote URLs for files, proof, video, and class access.")
    add_bullet(doc, "SI-3: The MVP shall not require a real payment gateway, file-storage API, SMS/email provider, or video SDK.")

    add_h3(doc, "3.3.4 Communication Interfaces")
    add_bullet(doc, "CI-1: React and Express shall communicate through RESTful HTTP/HTTPS endpoints using JSON.")
    add_bullet(doc, "CI-2: Protected requests shall use bearer JWTs and consistent success or error envelopes.")

    add_h1(doc, "4. Technology Stack & Architectural Overview")
    add_h2(doc, "4.1 Technology Stack Components")
    stack = [
        ("React.js", ": Responsive single-page presentation layer."),
        ("React Router, Axios & Vite", ": Routing, API communication, development, and bundling."),
        ("Node.js & Express.js", ": REST API, validation, authorization, and business workflows."),
        ("MySQL & mysql2", ": Relational persistence, constraints, search, transactions, and pooling."),
        ("JWT, bcryptjs, Helmet & CORS", ": Session identity, password hashing, headers, and origin control."),
        ("Node Test Runner, Supertest & Postman", ": Automated and manual API verification."),
    ]
    for number, (label, text) in enumerate(stack, 1):
        add_numbered(doc, number, text, label=label)

    add_h2(doc, "4.2 High-Level Architecture")
    architecture = [
        ("Presentation Layer (React.js)", ": Displays public and protected pages, handles forms, and calls the API."),
        ("Business Logic Layer (Express.js/Node.js)", ": Enforces validation, authorization, workflows, and responses."),
        ("Data Access Layer (Models/Query Builders)", ": Encapsulates persistence, search, aggregation, and transactions."),
        ("Data Layer (MySQL)", ": Stores identities, marketplace records, learning data, finance, trust, and engagement."),
    ]
    for number, (label, text) in enumerate(architecture, 1):
        add_numbered(doc, number, text, label=label)

    # Physical page 9
    add_h1(doc, "5. Tentative Development Plan (Agile Methodology)", page_break=True)
    sprint_content = [
        (
            "Sprint 1 (Weeks 1–2): Foundation, Authentication & Profiles",
            [
                "Establish the frontend, backend, environment configuration, MySQL schema, and seed data.",
                "Implement registration, login, JWT/RBAC, profiles, protected routes, and role layouts.",
            ],
        ),
        (
            "Sprint 2 (Weeks 3–4): Marketplace & Discovery",
            [
                "Implement tutor services, student requests, public details, search, filters, sorting, and pagination.",
                "Add saved tutors/courses, recent views, recommendation ordering, and course comparison.",
            ],
        ),
        (
            "Sprint 3 (Weeks 5–6): Applications, Bookings & Communication",
            [
                "Implement proposals, decisions, availability, trials, schedule conflicts, and booking transitions.",
                "Add persistent messages, unread state, flags, and in-app notifications.",
            ],
        ),
        (
            "Sprint 4 (Week 7): Learning, Finance & Trust Services",
            [
                "Implement materials, assignments, grading, quizzes, progress, reviews, and verification.",
                "Add the simulated payment ledger, commission, earnings, withdrawals, and safety reports.",
            ],
        ),
        (
            "Sprint 5 (Weeks 8–9): Administration, Quality & Delivery",
            [
                "Complete analytics, marketplace management, user status, and moderation workflows.",
                "Perform responsive, accessibility, security, testing, documentation, and final-demo work.",
            ],
        ),
    ]
    for heading, bullets in sprint_content:
        add_h3(doc, heading)
        for bullet in bullets:
            add_bullet(doc, bullet)

    # Physical page 10
    add_h1(doc, "6. Acceptance Criteria", page_break=True)
    criteria = [
        "The documented schema and seed data initialize in MySQL 8 or the compatible local MariaDB environment.",
        "The frontend, API, and database start through the documented npm commands; the health endpoint succeeds.",
        "Student, tutor, and provisioned administrator users authenticate and remain restricted to authorized operations.",
        "Visitors browse public tutors, services, reviews, and published student requests without authentication.",
        "A student can discover a tutor, publish a request, decide a proposal, create a conflict-free booking, learn, and make a mock payment.",
        "A tutor can publish a service, apply, manage bookings, teach, assess learners, view earnings, and submit verification evidence.",
        "The system prevents duplicate proposals, competing acceptances, booking conflicts, repeated trials, premature completion, duplicate payments, and invalid reviews.",
        "Protected materials, assignments, quizzes, and student records require the applicable teaching relationship.",
        "Finance calculations are server-derived and clearly identified as simulations rather than real settlement.",
        "An administrator can inspect analytics, verify tutors, moderate reports, manage user status, and process withdrawal records through protected operations.",
        "Backend automated tests pass and the React production build completes without errors.",
        "Main public and role workflows remain usable at common mobile and desktop viewport sizes.",
    ]
    for number, criterion in enumerate(criteria, 1):
        add_numbered(doc, number, f" {criterion}")

    add_h1(doc, "7. Conclusion")
    add_body(
        doc,
        "This SRS provides a detailed outline for the design and development of “Mentor Market”—a comprehensive tutoring marketplace using React.js, Node.js, Express.js, and MySQL. It integrates discovery, requests, applications, booking, communication, learning, simulated finance, trust, and administration while clearly separating MVP behavior from future real-world integrations. Following the outlined Agile sprints will support a robust, scalable, and user-friendly academic project.",
        after=0,
    )


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_sprint_heading(doc, text, color=BLACK, page_break=False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1
    set_run_font(
        p.add_run(text),
        name=ARIAL if color == BLACK else TIMES,
        size=20 if color == BLACK else 24,
        bold=color != BLACK,
        color=color,
    )
    return p


def add_feature_heading(doc, text, color=BLACK, size=None, highlight=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1
    if highlight:
        shade_paragraph(p, highlight)
    set_run_font(
        p.add_run(text),
        name=TIMES,
        size=size or (16 if color == BLACK else 24),
        bold=True,
        color=color,
    )
    return p


def set_picture_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_picture_fit(doc, filename, max_width=6.5, max_height=7.55, alt_text=None):
    path = ASSET_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"Required SRS screenshot not found: {path}")
    with Image.open(path) as image:
        width_px, height_px = image.size
    width = max_width
    height = width * height_px / width_px
    if height > max_height:
        height = max_height
        width = height * width_px / height_px
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    set_picture_alt_text(shape, alt_text or filename.replace("-", " ").replace(".png", ""))
    return p


def add_two_pictures(doc, first, second, max_height=3.25):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    for index, filename in enumerate((first, second)):
        cell = table.cell(index, 0)
        set_cell_margins(cell, top=25, start=0, bottom=25, end=0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = ASSET_DIR / filename
        if not path.exists():
            raise FileNotFoundError(f"Required SRS screenshot not found: {path}")
        with Image.open(path) as image:
            width_px, height_px = image.size
        width = min(6.5, max_height * width_px / height_px)
        height = width * height_px / width_px
        shape = p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
        set_picture_alt_text(shape, filename.replace("-", " ").replace(".png", ""))
    return table


APPENDIX_PAGES = [
    {
        "sprint": ("SPRINT 1", BLACK),
        "title": ("Feature 1 & 2: Registration and Role Authentication", BLACK, 16, None),
        "prose": "Visitors choose a student or tutor account, provide validated information, and create a role profile.",
        "images": ["register.png"],
    },
    {
        "title": ("Feature 3: Login", BLACK, 16, None),
        "prose": "Active users enter their credentials and are redirected to the workspace associated with their role.",
        "images": ["login.png"],
    },
    {
        "title": ("Feature 4: Public Home and Marketplace Entry", BLACK, 16, None),
        "images": ["home.png"],
    },
    {
        "title": ("Feature 5: Tutor Directory", BLACK, 16, None),
        "prose": "Visitors and students browse active tutor profiles and apply subject, mode, price, rating, location, and day filters.",
        "images": ["tutors.png"],
    },
    {
        "title": ("Feature 6: Student Dashboard", BLACK, 16, None),
        "images": ["student-dashboard.png"],
    },
    {
        "sprint": ("SPRINT 2", BLACK),
        "title": ("Feature 7: Personalized Course Discovery", BLACK, 16, None),
        "images": ["student-discovery.png"],
    },
    {
        "title": ("Feature 8: Course Details and Learning Path", BLACK, 16, None),
        "images": ["student-course.png"],
    },
    {
        "title": ("Feature 9: Saved Courses and Comparison", BLACK, 16, None),
        "images": ["student-saved-courses.png"],
    },
    {
        "title": ("Feature 10: Tutoring Request Publishing", BLACK, 16, None),
        "images": ["student-create-request.png"],
    },
    {
        "title": ("Feature 11: Applications and Matching", FEATURE_BLUE, 24, None),
        "images": ["student-applications.png"],
    },
    {
        "sprint": ("SPRINT 3", FEATURE_BLUE),
        "title": ("Feature 12: Booking Management", FEATURE_BLUE, 24, None),
        "images": ["student-bookings.png"],
    },
    {
        "title": ("Feature 13: Persistent Messaging", FEATURE_BLUE, 24, None),
        "images": ["student-messages.png"],
    },
    {
        "title": ("Feature 14: Study Materials", FEATURE_BLUE, 24, None),
        "images": ["student-materials.png"],
    },
    {
        "title": ("Feature 15: Assignments and Grading", FEATURE_BLUE, 24, None),
        "images": ["student-assignments.png"],
    },
    {
        "title": ("Feature 16: Quizzes and Automatic Scoring", FEATURE_BLUE, 24, RED),
        "images": ["student-quizzes.png"],
    },
    {
        "title": ("Feature 17: Progress Analytics", FEATURE_BLUE, 24, None),
        "images": ["student-progress.png"],
    },
    {
        "sprint": ("Sprint 4", FEATURE_BLUE),
        "title": ("Enhanced tutor workspace: improvement of earlier marketplace features", FEATURE_BLUE, 14, None),
        "images": ["tutor-dashboard.png"],
    },
    {
        "title": ("Feature 18: Tutor Creator Studio", FEATURE_BLUE, 24, None),
        "images": ["tutor-create-service.png"],
    },
    {
        "title": ("Service Listing Management", FEATURE_BLUE, 16, None),
        "images": ["tutor-services.png"],
    },
    {
        "title": ("Student Request Marketplace", FEATURE_BLUE, 16, None),
        "images": ["tutor-requests.png"],
    },
    {
        "title": ("Feature 19: Tutor Verification", FEATURE_BLUE, 24, RED),
        "images": ["tutor-verification.png"],
    },
    {
        "title": ("Feature 20: Earnings and Withdrawal Requests", FEATURE_BLUE, 24, None),
        "images": ["tutor-earnings.png"],
    },
    {
        "sprint": ("SPRINT 5", FEATURE_BLUE),
        "title": ("Feature 21: Simulated Payment History", FEATURE_BLUE, 24, None),
        "images": ["student-payments.png"],
    },
    {
        "title": ("Feature 22: Administrator Analytics", FEATURE_BLUE, 24, None),
        "images": ["admin-dashboard.png"],
    },
    {
        "title": ("Enhanced Admin Dashboard: user search and moderation", FEATURE_BLUE, 16, RED),
        "images": ["admin-users.png"],
    },
    {
        "title": ("Report Review and Marketplace Safety", FEATURE_BLUE, 16, DEEP_RED),
        "images": ["admin-reports.png"],
    },
]


def add_sprint_appendix(doc):
    for page_index, page in enumerate(APPENDIX_PAGES):
        sprint = page.get("sprint")
        if sprint:
            add_sprint_heading(
                doc,
                sprint[0],
                color=sprint[1],
                page_break=True,
            )
        else:
            title_probe = doc.add_paragraph()
            title_probe.paragraph_format.page_break_before = True
            title_probe.paragraph_format.space_after = Pt(0)
            title_probe.paragraph_format.line_spacing = 1
            # An empty page-start paragraph mirrors the pasted-evidence spacing in the reference.
            set_run_font(title_probe.add_run(" "), size=1, color=WHITE)

        title, color, size, highlight = page["title"]
        add_feature_heading(doc, title, color=color, size=size, highlight=highlight)
        if page.get("prose"):
            add_body(doc, page["prose"], after=11, align=WD_ALIGN_PARAGRAPH.LEFT)

        images = page["images"]
        if len(images) == 1:
            add_picture_fit(
                doc,
                images[0],
                max_width=6.5,
                max_height=7.45 if sprint else 7.9,
                alt_text=title,
            )
        else:
            add_two_pictures(doc, images[0], images[1])


def set_document_properties(doc):
    props = doc.core_properties
    props.title = "Mentor Market - CSE470 Project SRS"
    props.subject = "Software Requirements Specification for Mentor Market"
    props.author = "Sieuti Zaman; Mumtasim Daiyan; Abir Saha; Samina Fairooz Urbi"
    props.last_modified_by = "Mentor Market Project Team"
    props.keywords = "Mentor Market, CSE470, SRS, tutoring marketplace, software requirements"
    props.comments = (
        "Created in the visual format of the supplied WeHeal CSE470 SRS and grounded in "
        "the implemented Mentor Market repository."
    )


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    missing_assets = sorted(
        {
            filename
            for page in APPENDIX_PAGES
            for filename in page["images"]
            if not (ASSET_DIR / filename).exists()
        }
    )
    if missing_assets:
        raise FileNotFoundError(
            "The screenshot appendix is incomplete. Missing: " + ", ".join(missing_assets)
        )

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_formal_srs(doc)
    add_sprint_appendix(doc)
    set_document_properties(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
