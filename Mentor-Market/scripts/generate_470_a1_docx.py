"""Generate the Mentor Market CSE470 assignment in the supplied 470_A1 style.

The reference document is a plain US Letter academic handout:
one-inch margins, 13 pt Times New Roman, bold black headings, solid-circle
bullets, and no cover, header, footer, or page numbering.
"""

from __future__ import annotations

from pathlib import Path
import html
import math
import re
import subprocess
import sys
from xml.etree import ElementTree as ET


sys.path.insert(0, "/tmp/mentor_docx_deps")

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "documentation"
ASSET_DIR = OUTPUT_DIR / "assignment_assets"
OUTPUT = OUTPUT_DIR / "470_A1_Mentor_Market.docx"
DRAWIO_SOURCE = OUTPUT_DIR / "Mentor_Market_Class_Diagram.drawio"
DRAWIO_SVG = ASSET_DIR / "mentor_market_class_diagram_from_drawio.svg"
DOMAIN_DIAGRAM = ASSET_DIR / "mentor_market_class_diagram_from_drawio.png"
MVC_DIAGRAM = ASSET_DIR / "mentor_market_selected_features_mvc_class_diagram.png"

TIMES = "Times New Roman"
ARIAL = "Arial"
BLACK = "000000"
BODY_SIZE = 13

SERIF_REGULAR = Path("/usr/share/fonts/liberation/LiberationSerif-Regular.ttf")
SERIF_BOLD = Path("/usr/share/fonts/liberation/LiberationSerif-Bold.ttf")
SANS_REGULAR = Path("/usr/share/fonts/liberation/LiberationSans-Regular.ttf")
SANS_BOLD = Path("/usr/share/fonts/liberation/LiberationSans-Bold.ttf")
SANS_BOLD_ITALIC = Path("/usr/share/fonts/liberation/LiberationSans-BoldItalic.ttf")


Q1_INTRO = (
    "A requirement may be a high-level statement of a service or system constraint, "
    "or a detailed description of a function that the software must perform. For "
    "Mentor Market, the requirements are written at both levels. Functional "
    "requirements describe the services the system shall provide, how it shall "
    "react to input, and how it shall behave in particular situations. "
    "Non-functional requirements describe properties and constraints that usually "
    "apply to the system as a whole."
)


FUNCTIONAL_USER_INTRO = (
    "The lecture notes distinguish functional user requirements from functional "
    "system requirements. At the user level, the following statements describe "
    "what visitors, students, tutors, and administrators expect to accomplish "
    "without specifying how those goals will be implemented."
)


FUNCTIONAL_USER_REQUIREMENTS = [
    "FUR1: Students, tutors, and administrators shall be able to enter a workspace that is appropriate to their role after signing in.",
    "FUR2: Administrators shall be able to supervise marketplace accounts, tutor posts, student requests, verifications, and reports.",
    "FUR3: Visitors and students shall be able to discover tutors and compare available teaching services by searching, filtering, and sorting the results.",
    "FUR4: Visitors shall be able to learn what Mentor Market offers, understand how it works, and find contact information without creating an account.",
    "FUR5: Students shall be able to review a teaching service, its tutor, its suggested learning path, and its available class times before requesting a class.",
    "FUR6: Tutors shall be able to publish and manage the teaching services they offer through the marketplace.",
]


FUNCTIONAL_SYSTEM_INTRO = (
    "These user goals are refined below into functional system requirements that "
    "describe the services in enough detail to be implemented and checked. Each "
    "statement is written as a full sentence and names the specific role involved "
    "instead of using the general word “user.”"
)


FUNCTIONAL_SYSTEM_REQUIREMENTS = [
    (
        "FSR1 — Role-Based Workspaces",
        [
            "FSR1.1: After a successful login, the system shall take a student to course discovery and shall take a tutor or administrator to the appropriate dashboard.",
            "FSR1.2: The system shall provide students, tutors, and administrators with navigation and actions that match their assigned roles.",
            "FSR1.3: The system shall prevent a person from opening another role's protected workspace or performing an action that is not authorized for that role.",
            "FSR1.4: Each workspace shall show a clear loading state while information is being retrieved and an understandable empty or error message when information cannot be shown.",
        ],
    ),
    (
        "FSR2 — Administration and Marketplace Moderation",
        [
            "FSR2.1: The administrator dashboard shall summarize users, tutor posts, student requests, applications, bookings, payments, pending verifications, and open reports.",
            "FSR2.2: The system shall allow administrators to view account records, search the records shown in the administration workspace, and suspend or reactivate an account.",
            "FSR2.3: The system shall not allow an administrator to suspend their own account.",
            "FSR2.4: The system shall allow administrators to inspect tutor posts, student requests, reviews, bookings, payments, verifications, and reports.",
            "FSR2.5: The system shall allow administrators to approve or reject tutor verification requests and to resolve submitted reports.",
            "FSR2.6: The system shall allow administrators to delete inappropriate tutor posts, student requests, and reviews after confirming the decision.",
            "FSR2.7: After an administrative action, the system shall reload the affected information and explain whether the action succeeded or failed.",
        ],
    ),
    (
        "FSR3 — Tutor and Course Discovery",
        [
            "FSR3.1: The system shall allow visitors to browse the public tutor directory and shall allow students to browse active teaching services in their course-discovery feed.",
            "FSR3.2: The system shall allow visitors and students to search using relevant tutor names, subjects, qualifications, service titles, and service descriptions.",
            "FSR3.3: The public tutor directory shall allow filtering by subject, teaching mode, location, hourly rate, rating, and available day.",
            "FSR3.4: The student course feed shall allow filtering by subject, teaching mode, trial availability, and maximum price.",
            "FSR3.5: The system shall allow results to be sorted by suitable choices such as recommended order, rating, price, experience, or newest entry, depending on the selected directory.",
            "FSR3.6: Each result shall show enough information to compare an option, including the tutor or service name, subject, teaching mode, price, and other relevant details.",
            "FSR3.7: The system shall divide large result sets into pages or progressively loaded batches instead of displaying every record at once.",
            "FSR3.8: When no result matches the selected criteria, the system shall show a helpful empty message and a way to reset the search or filters.",
        ],
    ),
    (
        "FSR4 — Public Information Pages",
        [
            "FSR4.1: The system shall allow visitors to open the Home, About, How It Works, Become a Tutor, and Contact pages without signing in.",
            "FSR4.2: The Home and About pages shall explain the purpose of Mentor Market and the people it is designed to serve.",
            "FSR4.3: The How It Works page shall explain how students discover tutors and how tutors offer their services.",
            "FSR4.4: The Contact page shall show the project's support email address, telephone number, and location.",
            "FSR4.5: The public pages shall use consistent navigation, branding, footer information, and calls to action.",
        ],
    ),
    (
        "FSR5 — Teaching-Service Details and Learning Path",
        [
            "FSR5.1: The system shall retrieve the requested teaching service by its unique identifier.",
            "FSR5.2: The service page shall show its title, subject, level, price, teaching mode, availability, description, tutor information, and available reviews.",
            "FSR5.3: The system shall present subject-based learning outcomes and a suggested learning path made up of ordered modules and lessons.",
            "FSR5.4: Students shall be able to expand or collapse each module so that the learning path remains easy to read.",
            "FSR5.5: When a demonstration video is available, the system shall provide normal playback controls while keeping the written service information visible.",
            "FSR5.6: The system shall show available calendar slots and shall allow a student to request an available class time.",
            "FSR5.7: If the requested service cannot be found, the system shall show a friendly message and provide a way back to course discovery.",
        ],
    ),
    (
        "FSR6 — Tutor Service Creation and Management",
        [
            "FSR6.1: The system shall allow only an authenticated tutor to create a teaching-service post.",
            "FSR6.2: The system shall allow a tutor to enter a title, subject, level, price, teaching mode, availability, and description for the service.",
            "FSR6.3: Before publishing a service post, the system shall require a title, subject, level, price, teaching mode, availability, and description.",
            "FSR6.4: The system shall allow a tutor to add an optional location, thumbnail URL, demonstration-video URL, and trial-class option.",
            "FSR6.5: The system shall show a live preview of the service card while the tutor prepares the post.",
            "FSR6.6: The system shall allow a tutor to publish, view, pause, reactivate, and delete their own service posts.",
            "FSR6.7: The system shall ask for confirmation before deleting a service post.",
            "FSR6.8: The system shall not allow one tutor to modify another tutor's post, while an administrator shall remain able to remove a post through the moderation workspace.",
        ],
    ),
]


NON_FUNCTIONAL_INTRO = (
    "Functional services alone do not determine whether a system is acceptable. "
    "Non-functional requirements describe system properties and constraints such "
    "as response time, reliability, development methods, standards, and "
    "interoperability. They often affect the whole system, and failure to satisfy "
    "an important one may make the system unusable. Following the classification "
    "in the lecture slides, the requirements are grouped as Product, "
    "Organisational, and External requirements."
)


NON_FUNCTIONAL_REQUIREMENTS = [
    (
        "Product Requirements",
        [
            "NFR1: During a five-minute test with 20 concurrent clients and at least 500 normal non-media API requests, the system shall return at least 95% of responses within two seconds, excluding network and external-media transfer time.",
            "NFR2: The system shall return no more than 100 catalog records in one API response and shall provide pagination or progressive loading for additional results.",
            "NFR3: The system shall enforce authentication and the role, account-status, and ownership rules that apply to each protected action.",
            "NFR4: The system shall reject missing or invalid submitted values before changing persistent data.",
            "NFR5: The system shall not expose password hashes, private contact information, or administrative information on public pages.",
            "NFR6: Application decisions, tutor-verification reviews, tutor-profile updates, and withdrawal processing shall save all related changes together or leave the previous valid state unchanged.",
            "NFR7: The system shall identify missing or invalid form fields, ask for confirmation before destructive actions, and show distinct loading, success, empty, and error states.",
            "NFR8: At viewport widths of 360, 768, and 1280 pixels, the system shall present its public pages, forms, catalog, and role-based workspaces without unintended page-level horizontal scrolling.",
            "NFR9: The system shall support keyboard interaction, show visible focus on interactive controls, associate form fields with labels, and provide text alternatives for informative images.",
            "NFR10: The system shall keep essential written tutor and service information available when an external image or demonstration video is unavailable.",
            "NFR11: The system shall store monetary values as fixed-precision decimal values and shall restrict record statuses to the values supported by each workflow.",
        ],
    ),
    (
        "Organisational Requirements",
        [
            "NFR12: The project shall use React with Vite for the frontend, Node.js and Express.js for the backend, and MySQL for persistent data.",
            "NFR13: The backend shall follow the Model–View–Controller pattern and shall keep routes, controllers, and models in separate modules.",
            "NFR14: The backend shall use parameterized SQL queries for values received from requests.",
            "NFR15: The team shall keep the setup instructions, database schema, API examples, and important automated tests up to date as the project changes.",
        ],
    ),
    (
        "External Requirements",
        [
            "NFR16: The frontend and backend shall exchange information through HTTP requests and JSON responses so that the web interface and independent API clients can interoperate.",
            "NFR17: Because Mentor Market is an academic prototype, its payment and withdrawal workflows shall create simulated records only and shall not transfer real money through an external financial service.",
        ],
    ),
]


Q2_EXPLANATIONS = [
    "A UML class diagram gives a static view of an object-oriented system. The lecture describes a class as a blueprint for objects that share the same attributes, operations, relationships, and semantics. Each class in the model is drawn as a rectangle with separate compartments for the class name, attributes, and operations. Attributes follow the form attributeName : Type, while operations describe the behavior that an object can perform. The visibility marks - and + identify private attributes and public operations in this model.",
    "For Mentor Market, I prepared an analysis-level model in Draw.io instead of copying every database table or JavaScript module. The abstract User class is generalized into Student, Tutor, and Admin, while the abstract Listing class is generalized into Course and MentoringService. Named associations connect the user roles to other concepts, aggregation connects CourseCatalog and Listing, and composition models the Course–LearningPath–LearningModule–Lesson structure. The dashed dependency shows that CourseCatalog uses SearchCriteria, and multiplicities such as 1, 0..*, and 1..* state how many instances may participate. This scope keeps the diagram readable while covering the features specified in Q1.",
]


Q3_INTRO = (
    "For Mentor Market, I selected the Model–View–Controller (MVC) pattern. "
    "The lecture slides explain that MVC separates presentation and interaction "
    "from data-handling logic, creating separation of concern (SOC) between the "
    "interface and data processing. This separation is valuable because students, "
    "tutors, and administrators interact with the same marketplace information "
    "through different screens."
)


Q3_MVC_COMPONENTS = [
    "The Model contains the data-related logic of the application. In Mentor Market, model modules represent users, tutor profiles, tutor posts, student requests, bookings, messages, payments, and other persistent records. They communicate with MySQL through SELECT, INSERT, UPDATE, and DELETE operations and return the result to the Controller.",
    "The View is the interface presented to the end user. React pages and components, together with their HTML structure and CSS, provide Mentor Market's public directory, course-detail pages, forms, and role-based workspaces. These Views render dynamic values returned by the backend. The same service post can therefore appear as a discovery card to a student, an item on the tutor's service-management page, and a moderation record to an administrator.",
    "The Controller handles input received through URLs and coordinates the server-side request. Express controllers process GET, POST, PUT, PATCH, and DELETE requests, call the appropriate Model, and prepare the response required by the View. Authentication, role checks, and validation middleware are applied before protected requests reach the Controller.",
    "The slides illustrate a server-rendered View in which a Controller passes values directly to an HTML page. Mentor Market uses a client-rendered React interface instead. Its Controllers return JSON over HTTP, and React uses that response to render the page; the delivery mechanism is different, but the Model, View, and Controller responsibilities remain separated.",
]


Q3_ROUTES = (
    "The lecture defines routes as URLs used to access a resource. Mentor Market "
    "applies that idea through Express routes. For example, GET "
    "/api/tutor-posts/:id retrieves one service post, POST /api/tutor-posts "
    "creates a post, PATCH /api/tutor-posts/:id updates it, and DELETE "
    "/api/tutor-posts/:id removes it. A route identifies the correct Controller "
    "and applies the required authentication, role, and validation middleware "
    "instead of containing the whole workflow itself."
)


Q3_MVC_REASONS = [
    "The slides recommend MVC when data can be viewed and manipulated in several ways or when future presentation requirements are not fully known. Both conditions apply to Mentor Market. Tutor profiles and teaching services appear in public pages, student discovery tools, tutor-management screens, and administrative moderation tables. Because the data-related logic is separated from its representation, a page can be redesigned without changing how the record is stored, and a query can be improved without rebuilding every interface that uses it.",
    "MVC also provides the advantages identified in the lecture material: parallel work, improved testability, and code reusability. Frontend development can continue around an agreed route and JSON response while backend work continues in the Controller and Model. The trade-off is additional files and some extra complexity. That cost would be unnecessary for a very small program, but it is reasonable for Mentor Market because the project contains three roles and several connected workflows.",
]


Q3_REQUEST_CYCLE = [
    "A tutor enters the service information in the React listing View and submits the form.",
    "Axios sends a POST request to /api/tutor-posts.",
    "The Express route applies authentication, tutor-role, and validation middleware, then passes the request to the createPost Controller.",
    "The Controller adds the authenticated tutor's identifier and calls the TutorPost Model.",
    "The Model performs a parameterized INSERT operation in MySQL and returns the saved post.",
    "The Controller sends a successful 201 JSON response. React then navigates to the tutor's service-management View, which requests the service list again and displays the newly saved post.",
]


Q3_ALTERNATIVE = [
    "If MVC were not the required pattern, I would consider a layered architecture as a practical alternative. Mentor Market could be divided into a Presentation Layer for the React interface, an API/Application Layer for routes and request handling, a Business Layer for booking, moderation, and search rules, and a Data-Access Layer for MySQL operations. Each layer would provide a clear service to the layer above it.",
    "A layered design would separate the business rules more strictly and could become useful if the system grew much larger. However, it would also require additional service and repository modules and would make each request pass through more levels. MVC is the more direct choice for the present project because it is required by the course and already matches the routes, controllers, models, and React Views in the codebase.",
]


def set_run_font(run, name=TIMES, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    for mapping in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.rFonts.set(qn(mapping), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def set_style_font(style, name=TIMES, size=BODY_SIZE, bold=False):
    style.font.name = name
    rpr = style._element.get_or_add_rPr()
    for mapping in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.rFonts.set(qn(mapping), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def set_page_geometry(section, landscape=False, diagram=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        margin = 0.15 if diagram else 1
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        margin = 1
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def configure_document(doc):
    set_page_geometry(doc.sections[0])
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        set_style_font(style, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    styles["Heading 1"].paragraph_format.space_before = Pt(0)

    for section in doc.sections:
        section.header.paragraphs[0].clear()
        section.footer.paragraphs[0].clear()

    compat = doc.settings._element.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        doc.settings._element.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)


def add_heading(doc, text, page_break=False, first=False):
    paragraph = doc.add_paragraph(style="Heading 2")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.space_before = Pt(0 if first else 12)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(text), bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(paragraph.add_run("●"), name=ARIAL)
    paragraph.add_run("\t")
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_body(doc, text, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(paragraph.add_run(text), italic=italic)
    return paragraph


def add_page_break(doc):
    return doc.add_page_break()


def _font(path: Path, size: int):
    return ImageFont.truetype(str(path), size)


def _fit_font(draw, text, path, preferred, max_width, minimum=22):
    size = preferred
    while size > minimum:
        font = _font(path, size)
        if draw.textbbox((0, 0), text, font=font)[2] <= max_width:
            return font
        size -= 1
    return _font(path, minimum)


def _draw_centered(draw, box, text, font, y, fill="black"):
    left, _, right, _ = box
    width = draw.textbbox((0, 0), text, font=font)[2]
    draw.text(((left + right - width) / 2, y), text, font=font, fill=fill)


def draw_class_box(draw, box, name, stereotype, attributes=None, methods=None):
    attributes = attributes or []
    methods = methods or []
    left, top, right, bottom = box
    draw.rectangle(box, outline="black", width=3, fill="white")

    stereo_font = _font(SERIF_REGULAR, 24)
    name_font = _fit_font(draw, name, SERIF_BOLD, 34, right - left - 24)
    body_font = _font(SERIF_REGULAR, 25)
    header_bottom = top + 82

    _draw_centered(draw, box, f"«{stereotype}»", stereo_font, top + 8)
    _draw_centered(draw, box, name, name_font, top + 39)
    draw.line((left, header_bottom, right, header_bottom), fill="black", width=2)

    line_height = 31
    cursor = header_bottom + 10
    for text in attributes:
        draw.text((left + 14, cursor), text, font=body_font, fill="black")
        cursor += line_height

    if methods:
        separator = cursor + 4
        draw.line((left, separator, right, separator), fill="black", width=2)
        cursor = separator + 9
        for text in methods:
            draw.text((left + 14, cursor), text, font=body_font, fill="black")
            cursor += line_height


def draw_arrow(draw, start, end, dashed=False, width=3):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        distance = math.hypot(x2 - x1, y2 - y1)
        if distance:
            dx = (x2 - x1) / distance
            dy = (y2 - y1) / distance
            segment = 14
            gap = 10
            position = 0
            while position < distance - 18:
                finish = min(position + segment, distance - 18)
                draw.line(
                    (x1 + dx * position, y1 + dy * position, x1 + dx * finish, y1 + dy * finish),
                    fill="black",
                    width=width,
                )
                position += segment + gap
    else:
        draw.line((x1, y1, x2, y2), fill="black", width=width)

    angle = math.atan2(y2 - y1, x2 - x1)
    length = 18
    spread = math.pi / 7
    p1 = (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread))
    p2 = (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread))
    draw.line((p1[0], p1[1], x2, y2, p2[0], p2[1]), fill="black", width=width)


def draw_hollow_triangle(draw, apex, direction="up", size=22):
    x, y = apex
    if direction == "up":
        points = [(x, y), (x - size, y + size * 1.5), (x + size, y + size * 1.5)]
    else:
        points = [(x, y), (x - size, y - size * 1.5), (x + size, y - size * 1.5)]
    draw.polygon(points, outline="black", fill="white")


def draw_diamond(draw, center, direction="right", size=17):
    x, y = center
    if direction in ("right", "left"):
        points = [(x - size, y), (x, y - size * 0.7), (x + size, y), (x, y + size * 0.7)]
    else:
        points = [(x, y - size), (x - size * 0.7, y), (x, y + size), (x + size * 0.7, y)]
    draw.polygon(points, outline="black", fill="black")


def generate_domain_diagram():
    width, height = 3000, 1800
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    relation_font = _font(SERIF_REGULAR, 27)

    summary_box = (60, 80, 790, 400)
    user_box = (1060, 60, 1940, 390)
    static_box = (2280, 60, 2940, 490)
    student_box = (560, 560, 1010, 820)
    teacher_box = (1270, 560, 1730, 820)
    admin_box = (1990, 560, 2450, 820)
    criteria_box = (60, 970, 790, 1370)
    course_box = (1000, 960, 1840, 1710)
    path_box = (2070, 920, 2940, 1165)
    module_box = (2070, 1280, 2940, 1515)
    lesson_box = (2070, 1630, 2940, 1780)

    draw_class_box(
        draw,
        user_box,
        "User",
        "Model",
        ["+Long id", "+String fullName", "+String email", "+String role", "+Boolean isActive"],
    )
    draw_class_box(draw, student_box, "Student", "Model", methods=["+openDashboard()", "+browseCatalog()", "+viewCourse()"])
    draw_class_box(draw, teacher_box, "Teacher", "Model", methods=["+openDashboard()", "+createCourse()", "+manageOwnCourses()"])
    draw_class_box(draw, admin_box, "Admin", "Model", methods=["+openDashboard()", "+setUserStatus()", "+moderateCourse()"])
    draw_class_box(
        draw,
        summary_box,
        "DashboardSummary",
        "Model",
        ["+String role", "+Map metrics", "+List quickActions"],
    )
    draw_class_box(
        draw,
        course_box,
        "CourseOffering",
        "Model",
        [
            "+Long id",
            "+Long teacherId",
            "+String title",
            "+String subject",
            "+String level",
            "+Decimal price",
            "+String teachingMode",
            "+String availability",
            "+String thumbnailUrl",
            "+String demoVideoUrl",
            "+String description",
            "+String status",
        ],
        ["+validate()", "+activate()", "+deactivate()"],
    )
    draw_class_box(
        draw,
        criteria_box,
        "SearchCriteria",
        "Model",
        ["+String keyword", "+String subject", "+String teachingMode", "+Decimal maxPrice", "+String sortBy", "+Integer page"],
    )
    draw_class_box(
        draw,
        path_box,
        "LearningPath",
        "Model",
        ["+String title", "+List outcomes"],
    )
    draw_class_box(
        draw,
        module_box,
        "CourseModule",
        "Model",
        ["+Integer sequence", "+String title", "+String summary"],
    )
    draw_class_box(
        draw,
        lesson_box,
        "Lesson",
        "Model",
        ["+Integer sequence", "+String title"],
    )
    draw_class_box(
        draw,
        static_box,
        "StaticPageContent",
        "Model",
        ["+String slug", "+String title", "+String body", "+String mediaUrl", "+List contentItems"],
        ["+getBySlug(slug)"],
    )

    # Inheritance: Student, Teacher, and Admin point to User.
    fork_y = 485
    parent_x = (user_box[0] + user_box[2]) / 2
    draw.line((parent_x, user_box[3], parent_x, fork_y), fill="black", width=3)
    draw.line(((student_box[0] + student_box[2]) / 2, fork_y, (admin_box[0] + admin_box[2]) / 2, fork_y), fill="black", width=3)
    for role_box in (student_box, teacher_box, admin_box):
        role_x = (role_box[0] + role_box[2]) / 2
        draw.line((role_x, fork_y, role_x, role_box[1]), fill="black", width=3)
    draw_hollow_triangle(draw, (parent_x, user_box[3]), direction="up")

    # Dashboard summary is built for a user role.
    draw_arrow(draw, (summary_box[2], 245), (user_box[0], 245), dashed=True)
    draw.text((820, 205), "describes role", font=relation_font, fill="black")

    # Ownership association.
    teacher_x = (teacher_box[0] + teacher_box[2]) / 2
    draw.line((teacher_x, teacher_box[3], teacher_x, course_box[1]), fill="black", width=3)
    draw.text((teacher_x + 22, 870), "1 owns 0..*", font=relation_font, fill="black")

    # Admin moderation authority.
    draw_arrow(draw, (admin_box[0], 760), (course_box[2], 1040), dashed=True)
    draw.text((1860, 875), "moderates", font=relation_font, fill="black")

    # Search criteria filters the course catalog.
    draw_arrow(draw, (criteria_box[2], 1110), (course_box[0], 1110), dashed=True)
    draw.text((820, 1068), "filters", font=relation_font, fill="black")

    # Course-learning-path composition.
    composition_y = 1040
    draw.line((course_box[2], composition_y, path_box[0], composition_y), fill="black", width=3)
    draw_diamond(draw, (course_box[2] + 18, composition_y), direction="right")
    draw.text((1900, composition_y - 38), "1", font=relation_font, fill="black")

    # Learning path contains modules; modules contain lessons.
    path_x = (path_box[0] + path_box[2]) / 2
    draw.line((path_x, path_box[3], path_x, module_box[1]), fill="black", width=3)
    draw_diamond(draw, (path_x, path_box[3] + 18), direction="down")
    draw.text((path_x + 20, 1200), "1..*", font=relation_font, fill="black")

    module_x = (module_box[0] + module_box[2]) / 2
    draw.line((module_x, module_box[3], module_x, lesson_box[1]), fill="black", width=3)
    draw_diamond(draw, (module_x, module_box[3] + 18), direction="down")
    draw.text((module_x + 20, 1560), "1..*", font=relation_font, fill="black")

    image.save(DOMAIN_DIAGRAM, dpi=(300, 300), optimize=True)


def generate_mvc_diagram():
    width, height = 3000, 1500
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    label_font = _font(SERIF_BOLD, 29)
    note_font = _font(SERIF_REGULAR, 25)

    draw.text((28, 16), "VIEWS", font=label_font, fill="black")
    draw.text((28, 290), "CONTROLLERS", font=label_font, fill="black")
    draw.text((28, 805), "PRIMARY MODELS", font=label_font, fill="black")

    x_positions = [40, 535, 1030, 1525, 2020, 2515]
    specs = [
        (
            "DashboardView",
            "DashboardController",
            ["+getDashboard(userId, role)"],
            "DashboardSummary",
            ["+String role", "+Map metrics", "+List quickActions"],
        ),
        (
            "AdminModerationView",
            "ModerationController",
            ["+listResources(criteria)", "+setUserStatus(...)", "+moderateCourse(...)"],
            "User",
            ["+Long id", "+String role", "+Boolean isActive"],
        ),
        (
            "CourseCatalogView",
            "CatalogController",
            ["+searchCourses(criteria)"],
            "SearchCriteria",
            ["+String keyword", "+String subject", "+String mode", "+Decimal maxPrice", "+String sortBy"],
        ),
        (
            "CourseDetailView",
            "CourseDetailController",
            ["+getCourse(courseId)", "+getLearningPath(courseId)"],
            "LearningPath",
            ["+String title", "+List outcomes", "+List modules"],
        ),
        (
            "CourseManagementView",
            "CourseManagementController",
            ["+listOwnCourses(teacherId)", "+createCourse(...)", "+updateCourse(...)", "+changeStatus(...)", "+deleteCourse(...)"],
            "CourseOffering",
            ["+Long id", "+Long teacherId", "+String title", "+Decimal price", "+String status"],
        ),
        (
            "PublicPageView",
            "PublicPageController",
            ["+getPage(slug)"],
            "StaticPageContent",
            ["+String slug", "+String title", "+String body", "+String mediaUrl"],
        ),
    ]

    for x, (view_name, controller_name, controller_methods, model_name, model_attributes) in zip(x_positions, specs):
        view_box = (x, 65, x + 445, 235)
        controller_box = (x, 345, x + 445, 745)
        model_box = (x, 870, x + 445, 1325)
        draw_class_box(draw, view_box, view_name, "View", methods=["+render(data)"])
        draw_class_box(draw, controller_box, controller_name, "Controller", methods=controller_methods)
        draw_class_box(draw, model_box, model_name, "Model", attributes=model_attributes)
        center = x + 222.5
        draw_arrow(draw, (center, view_box[3]), (center, controller_box[1]), dashed=True)
        draw_arrow(draw, (center, controller_box[3]), (center, model_box[1]), dashed=True)

    draw.text(
        (40, 1405),
        "Shared domain relationships and secondary model dependencies are detailed in the Domain Class Diagram.",
        font=note_font,
        fill="black",
    )
    image.save(MVC_DIAGRAM, dpi=(300, 300), optimize=True)


def _fit_lines_font(draw, lines, path, preferred, max_width, minimum=18):
    lines = [line for line in lines if line]
    if not lines:
        return _font(path, preferred)
    size = preferred
    while size > minimum:
        font = _font(path, size)
        if max(draw.textbbox((0, 0), line, font=font)[2] for line in lines) <= max_width:
            return font
        size -= 1
    return _font(path, minimum)


def _draw_label(draw, x, y, text, font, padding=5):
    bounds = draw.textbbox((0, 0), text, font=font)
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    box = (
        x - width / 2 - padding,
        y - height / 2 - padding,
        x + width / 2 + padding,
        y + height / 2 + padding,
    )
    draw.rectangle(box, fill="white")
    draw.text((x - width / 2, y - height / 2 - bounds[1]), text, font=font, fill="black")


def draw_uml_class(draw, box, name, attributes=None, operations=None, abstract=False):
    """Draw one reference-style UML class with conventional compartments."""
    attributes = attributes or []
    operations = operations or []
    left, top, right, bottom = box
    draw.rectangle(box, outline="black", width=3, fill="white")

    if abstract:
        stereotype_font = _font(SANS_REGULAR, 21)
        name_font = _fit_font(draw, name, SANS_BOLD_ITALIC, 35, right - left - 28)
        stereotype = "«abstract»"
        stereotype_width = draw.textbbox((0, 0), stereotype, font=stereotype_font)[2]
        name_width = draw.textbbox((0, 0), name, font=name_font)[2]
        draw.text(((left + right - stereotype_width) / 2, top + 8), stereotype, font=stereotype_font, fill="black")
        draw.text(((left + right - name_width) / 2, top + 37), name, font=name_font, fill="black")
        header_bottom = top + 86
    else:
        name_font = _fit_font(draw, name, SANS_BOLD, 35, right - left - 28)
        name_width = draw.textbbox((0, 0), name, font=name_font)[2]
        draw.text(((left + right - name_width) / 2, top + 18), name, font=name_font, fill="black")
        header_bottom = top + 67

    draw.line((left, header_bottom, right, header_bottom), fill="black", width=2)

    members = [*attributes, *operations]
    body_height = bottom - header_bottom - 24 - (16 if attributes and operations else 0)
    height_size = max(18, int(body_height / max(1, len(members))) - 6)
    preferred_size = min(38, height_size)
    body_font = _fit_lines_font(draw, members, SANS_REGULAR, preferred_size, right - left - 28)
    line_height = body_font.size + 7
    cursor = header_bottom + 13

    for text in attributes:
        draw.text((left + 14, cursor), text, font=body_font, fill="black")
        cursor += line_height

    if attributes and operations:
        separator = cursor + 4
        draw.line((left, separator, right, separator), fill="black", width=2)
        cursor = separator + 11

    for text in operations:
        draw.text((left + 14, cursor), text, font=body_font, fill="black")
        cursor += line_height


def draw_uml_triangle(draw, apex, size=25):
    x, y = apex
    points = [(x, y), (x - size, y + size * 1.5), (x + size, y + size * 1.5)]
    draw.polygon(points, fill="white")
    draw.line((*points[0], *points[1], *points[2], *points[0]), fill="black", width=3)
    return y + size * 1.5


def draw_uml_diamond(draw, endpoint, direction, filled):
    """Draw a UML diamond whose first vertex touches the whole class."""
    x, y = endpoint
    length = 35
    half_height = 13
    if direction == "right":
        points = [(x, y), (x + length / 2, y - half_height), (x + length, y), (x + length / 2, y + half_height)]
        far = (x + length, y)
    elif direction == "down":
        points = [(x, y), (x - half_height, y + length / 2), (x, y + length), (x + half_height, y + length / 2)]
        far = (x, y + length)
    else:
        raise ValueError(f"Unsupported diamond direction: {direction}")
    draw.polygon(points, outline="black", fill="black" if filled else "white")
    draw.line((*points[0], *points[1], *points[2], *points[3], *points[0]), fill="black", width=3)
    return far


def draw_dashed_vertical_dependency(draw, start, end):
    x, y1 = start
    _, y2 = end
    position = y1
    while position < y2 - 24:
        finish = min(position + 15, y2 - 24)
        draw.line((x, position, x, finish), fill="black", width=3)
        position += 25
    draw.line((x - 12, y2 - 19, x, y2, x + 12, y2 - 19), fill="black", width=3)


def generate_reference_class_diagram():
    """Generate one domain UML diagram matching the supplied CSE470 lecture."""
    width, height = 3300, 2200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    relationship_font = _font(SANS_REGULAR, 24)
    multiplicity_font = _font(SANS_REGULAR, 23)

    dashboard_box = (30, 40, 850, 410)
    user_box = (1130, 30, 2170, 460)
    public_page_box = (2440, 40, 3270, 460)
    student_box = (80, 590, 850, 920)
    teacher_box = (1260, 590, 2040, 1010)
    admin_box = (2380, 590, 3260, 1010)
    catalog_box = (30, 1120, 850, 1500)
    criteria_box = (30, 1700, 850, 2160)
    listing_box = (1100, 1100, 2200, 1810)
    service_box = (1060, 1930, 1580, 2190)
    course_box = (1660, 1930, 2200, 2190)
    path_box = (2410, 1080, 3270, 1370)
    module_box = (2410, 1500, 3270, 1810)
    lesson_box = (2410, 1930, 3270, 2190)

    draw_uml_class(
        draw,
        dashboard_box,
        "Dashboard",
        ["- dashboardType : String", "- lastUpdatedAt : DateTime"],
        ["+ loadFor(userId : Long)", "+ displaySummary()"],
    )
    draw_uml_class(
        draw,
        user_box,
        "User",
        [
            "- userId : Long",
            "- fullName : String",
            "- email : String",
            "- accountStatus : String",
        ],
        [
            "+ login(email : String, password : String)",
            "+ logout()",
            "+ openDashboard()",
        ],
        abstract=True,
    )
    draw_uml_class(
        draw,
        public_page_box,
        "PublicPage",
        [
            "- pageId : Long",
            "- pageType : PageType",
            "- title : String",
            "- body : String",
            "- supportEmail : String [0..1]",
            "- supportPhone : String [0..1]",
        ],
        ["+ getByType(type : PageType)"],
    )
    draw_uml_class(
        draw,
        student_box,
        "Student",
        operations=[
            "+ browseCatalog(criteria : SearchCriteria)",
            "+ viewCourse(courseId : Long)",
        ],
    )
    draw_uml_class(
        draw,
        teacher_box,
        "Teacher",
        operations=[
            "+ createListing(data : Listing)",
            "+ updateListing(id : Long)",
            "+ publishListing(id : Long)",
            "+ setStatus(id : Long, status : String)",
            "+ deleteListing(id : Long)",
        ],
    )
    draw_uml_class(
        draw,
        admin_box,
        "Admin",
        operations=[
            "+ viewMarketplace()",
            "+ suspendUser(id : Long)",
            "+ reactivateUser(id : Long)",
            "+ moderateListing(id : Long, status : String)",
            "+ deleteListing(id : Long)",
        ],
    )
    draw_uml_class(
        draw,
        catalog_box,
        "CourseCatalog",
        ["- pageSize : Integer", "- currentPage : Integer"],
        [
            "+ search(criteria : SearchCriteria)",
            "+ filter(criteria : SearchCriteria)",
            "+ sort(option : SortOption)",
            "+ resetFilters()",
        ],
    )
    draw_uml_class(
        draw,
        criteria_box,
        "SearchCriteria",
        [
            "- keyword : String [0..1]",
            "- subject : String [0..1]",
            "- teachingMode : String [0..1]",
            "- minPrice : Decimal [0..1]",
            "- maxPrice : Decimal [0..1]",
            "- sortBy : SortOption",
            "- page : Integer",
        ],
    )
    draw_uml_class(
        draw,
        listing_box,
        "Listing",
        [
            "- listingId : Long",
            "- title : String",
            "- subject : String",
            "- level : String",
            "- price : Decimal",
            "- teachingMode : String",
            "- availability : String",
            "- status : ListingStatus",
        ],
        [
            "+ validate()",
            "+ publish()",
            "+ updateDetails()",
            "+ setActive(active : Boolean)",
        ],
        abstract=True,
    )
    draw_uml_class(
        draw,
        service_box,
        "MentoringService",
        ["- serviceType : String"],
        ["+ getServiceDetails()"],
    )
    draw_uml_class(
        draw,
        course_box,
        "Course",
        ["- learningOutcomes : List<String>"],
        ["+ getLearningPath()"],
    )
    draw_uml_class(
        draw,
        path_box,
        "LearningPath",
        ["- pathId : Long", "- title : String", "- outcomes : List<String>"],
        ["+ getOrderedModules()"],
    )
    draw_uml_class(
        draw,
        module_box,
        "LearningModule",
        ["- moduleId : Long", "- sequenceNo : Integer", "- title : String", "- summary : String"],
        ["+ addLesson(lesson : Lesson)"],
    )
    draw_uml_class(
        draw,
        lesson_box,
        "Lesson",
        ["- lessonId : Long", "- sequenceNo : Integer", "- title : String"],
        ["+ open()"],
    )

    # User role generalization: subclasses share one hollow triangle at User.
    user_center = (user_box[0] + user_box[2]) / 2
    fork_y = 530
    triangle_base = draw_uml_triangle(draw, (user_center, user_box[3]))
    draw.line((user_center, triangle_base, user_center, fork_y), fill="black", width=3)
    role_centers = [(box[0] + box[2]) / 2 for box in (student_box, teacher_box, admin_box)]
    draw.line((role_centers[0], fork_y, role_centers[-1], fork_y), fill="black", width=3)
    for center, box in zip(role_centers, (student_box, teacher_box, admin_box)):
        draw.line((center, fork_y, center, box[1]), fill="black", width=3)

    # One User accesses one role-specific Dashboard.
    association_y = 245
    draw.line((dashboard_box[2], association_y, user_box[0], association_y), fill="black", width=3)
    _draw_label(draw, 990, 216, "accesses", relationship_font)
    draw.text((865, 250), "1", font=multiplicity_font, fill="black")
    draw.text((1090, 250), "1", font=multiplicity_font, fill="black")

    # Users can view the four publicly accessible page types.
    draw.line((user_box[2], association_y, public_page_box[0], association_y), fill="black", width=3)
    _draw_label(draw, 2305, 216, "views", relationship_font)
    draw.text((2180, 250), "0..*", font=multiplicity_font, fill="black")
    draw.text((2380, 250), "0..*", font=multiplicity_font, fill="black")

    # Student browses the shared catalog.
    student_center = (student_box[0] + student_box[2]) / 2
    draw.line((student_center, student_box[3], student_center, catalog_box[1]), fill="black", width=3)
    _draw_label(draw, student_center + 68, 1015, "browses", relationship_font)
    draw.text((student_center + 14, 940), "0..*", font=multiplicity_font, fill="black")
    draw.text((student_center + 14, 1085), "1", font=multiplicity_font, fill="black")

    # A Teacher owns and manages many marketplace listings.
    teacher_center = (teacher_box[0] + teacher_box[2]) / 2
    draw.line((teacher_center, teacher_box[3], teacher_center, listing_box[1]), fill="black", width=3)
    _draw_label(draw, teacher_center + 115, 1052, "owns / manages", relationship_font)
    draw.text((teacher_center + 14, 1020), "1", font=multiplicity_font, fill="black")
    draw.text((teacher_center + 14, 1070), "0..*", font=multiplicity_font, fill="black")

    # Admins moderate marketplace listings.
    admin_start = (admin_box[0], 840)
    admin_route_x = 2290
    listing_endpoint = (listing_box[2], 1460)
    draw.line(
        (
            admin_start[0],
            admin_start[1],
            admin_route_x,
            admin_start[1],
            admin_route_x,
            listing_endpoint[1],
            listing_endpoint[0],
            listing_endpoint[1],
        ),
        fill="black",
        width=3,
        joint="curve",
    )
    _draw_label(draw, admin_route_x, 1090, "moderates", relationship_font)
    draw.text((2330, 802), "0..*", font=multiplicity_font, fill="black")
    draw.text((2215, 1424), "0..*", font=multiplicity_font, fill="black")

    # CourseCatalog aggregates active marketplace listings.
    catalog_y = 1300
    diamond_far = draw_uml_diamond(draw, (catalog_box[2], catalog_y), "right", filled=False)
    draw.line((*diamond_far, listing_box[0], catalog_y), fill="black", width=3)
    _draw_label(draw, 978, 1263, "lists", relationship_font)
    draw.text((860, 1315), "1", font=multiplicity_font, fill="black")
    draw.text((1040, 1315), "0..*", font=multiplicity_font, fill="black")

    # Search criteria is a value object used by the catalog.
    criteria_x = (catalog_box[0] + catalog_box[2]) / 2
    draw_dashed_vertical_dependency(draw, (criteria_x, catalog_box[3]), (criteria_x, criteria_box[1]))
    _draw_label(draw, criteria_x + 48, 1600, "uses", relationship_font)

    # Course and MentoringService specialize MarketplaceListing.
    listing_center = (listing_box[0] + listing_box[2]) / 2
    listing_triangle_base = draw_uml_triangle(draw, (listing_center, listing_box[3]))
    listing_fork_y = 1855
    draw.line((listing_center, listing_triangle_base, listing_center, listing_fork_y), fill="black", width=3)
    subtype_centers = [(service_box[0] + service_box[2]) / 2, (course_box[0] + course_box[2]) / 2]
    draw.line((subtype_centers[0], listing_fork_y, subtype_centers[1], listing_fork_y), fill="black", width=3)
    for center, box in zip(subtype_centers, (service_box, course_box)):
        draw.line((center, listing_fork_y, center, box[1]), fill="black", width=3)

    # A Course owns its learning path; paths own modules; modules own lessons.
    course_y = 2045
    course_diamond_far = draw_uml_diamond(draw, (course_box[2], course_y), "right", filled=True)
    course_route_x = 2330
    path_y = 1280
    draw.line(
        (
            course_diamond_far[0],
            course_y,
            course_route_x,
            course_y,
            course_route_x,
            path_y,
            path_box[0],
            path_y,
        ),
        fill="black",
        width=3,
        joint="curve",
    )
    _draw_label(draw, course_route_x - 12, 1885, "defines", relationship_font)
    draw.text((2210, 2060), "1", font=multiplicity_font, fill="black")
    draw.text((2370, 1225), "1", font=multiplicity_font, fill="black")

    chain_x = (path_box[0] + path_box[2]) / 2
    path_diamond_far = draw_uml_diamond(draw, (chain_x, path_box[3]), "down", filled=True)
    draw.line((*path_diamond_far, chain_x, module_box[1]), fill="black", width=3)
    _draw_label(draw, chain_x + 155, 1480, "contains", relationship_font)
    draw.text((chain_x + 18, 1420), "1", font=multiplicity_font, fill="black")
    draw.text((chain_x + 18, 1463), "1..*", font=multiplicity_font, fill="black")

    module_diamond_far = draw_uml_diamond(draw, (chain_x, module_box[3]), "down", filled=True)
    draw.line((*module_diamond_far, chain_x, lesson_box[1]), fill="black", width=3)
    _draw_label(draw, chain_x + 155, 1915, "contains", relationship_font)
    draw.text((chain_x + 18, 1830), "1", font=multiplicity_font, fill="black")
    draw.text((chain_x + 18, 1893), "1..*", font=multiplicity_font, fill="black")

    image.save(DOMAIN_DIAGRAM, dpi=(300, 300), optimize=True)


def _mx_style(style_text):
    """Turn a diagrams.net style string into a small key/value mapping."""
    result = {}
    for item in (style_text or "").split(";"):
        if not item:
            continue
        if "=" in item:
            key, value = item.split("=", 1)
            result[key] = value
        else:
            result[item] = "1"
    return result


def _mx_geometry(cell):
    geometry = cell.find("mxGeometry")
    if geometry is None:
        return (0.0, 0.0, 0.0, 0.0)

    def number(name):
        try:
            return float(geometry.get(name, "0"))
        except ValueError:
            return 0.0

    return (number("x"), number("y"), number("width"), number("height"))


def _drawio_text_parts(value):
    """Read the simple HTML labels used by this Draw.io class diagram."""
    raw = value or ""
    raw = re.sub(r"<hr\s*/?>", "\n__COMPARTMENT__\n", raw, flags=re.I)
    raw = re.sub(r"<br\s*/?>", "\n", raw, flags=re.I)
    raw = re.sub(r"</div\s*>", "\n", raw, flags=re.I)
    raw = re.sub(r"<[^>]+>", "", raw)
    raw = html.unescape(raw)

    parts = []
    for part in raw.split("__COMPARTMENT__"):
        lines = [line.strip() for line in part.splitlines() if line.strip()]
        if lines:
            parts.append(lines)
    return parts


def _svg_text(x, y, text, size, anchor="start", weight="normal", style="normal"):
    return (
        f'<text x="{x:.2f}" y="{y:.2f}" font-family="Arial, Liberation Sans, sans-serif" '
        f'font-size="{size}" font-weight="{weight}" font-style="{style}" '
        f'text-anchor="{anchor}" fill="#000000">{html.escape(text)}</text>'
    )


def _cell_anchor(bounds, style, prefix, other_center):
    x, y, width, height = bounds
    x_key = f"{prefix}X"
    y_key = f"{prefix}Y"
    if x_key in style and y_key in style:
        return (
            x + float(style[x_key]) * width,
            y + float(style[y_key]) * height,
        )
    if width <= 3 and height <= 3:
        return (x + width / 2, y + height / 2)

    center = (x + width / 2, y + height / 2)
    dx = other_center[0] - center[0]
    dy = other_center[1] - center[1]
    if abs(dx) > abs(dy):
        return (x + width if dx > 0 else x, center[1])
    return (center[0], y + height if dy > 0 else y)


def _orthogonal_points(start, end, waypoints):
    if waypoints:
        points = [start, *waypoints, end]
    elif abs(start[0] - end[0]) < 0.1 or abs(start[1] - end[1]) < 0.1:
        points = [start, end]
    elif abs(start[0] - end[0]) > abs(start[1] - end[1]):
        middle_x = (start[0] + end[0]) / 2
        points = [start, (middle_x, start[1]), (middle_x, end[1]), end]
    else:
        middle_y = (start[1] + end[1]) / 2
        points = [start, (start[0], middle_y), (end[0], middle_y), end]

    cleaned = [points[0]]
    for point in points[1:]:
        if point != cleaned[-1]:
            cleaned.append(point)
    return cleaned


def _edge_label_position(points):
    segments = []
    for start, end in zip(points, points[1:]):
        length = abs(end[0] - start[0]) + abs(end[1] - start[1])
        segments.append((length, start, end))
    _, start, end = max(segments, key=lambda item: item[0])
    return (
        (start[0] + end[0]) / 2,
        (start[1] + end[1]) / 2,
        abs(start[0] - end[0]) < abs(start[1] - end[1]),
    )


def export_drawio_class_diagram():
    """Render the editable Draw.io XML directly into SVG and PNG.

    The Draw.io file remains the canonical source. This renderer reads its
    class labels, coordinates, relationship endpoints, waypoints, arrows,
    diamonds, and multiplicities instead of recreating the model from a
    separate hard-coded diagram definition.
    """
    if not DRAWIO_SOURCE.exists():
        raise FileNotFoundError(f"Draw.io source not found: {DRAWIO_SOURCE}")

    tree = ET.parse(DRAWIO_SOURCE)
    model = tree.find("./diagram/mxGraphModel")
    if model is None:
        raise ValueError("The Draw.io file does not contain an mxGraphModel")
    graph_root = model.find("root")
    if graph_root is None:
        raise ValueError("The Draw.io file does not contain a graph root")

    page_width = float(model.get("pageWidth", "1900"))
    page_height = float(model.get("pageHeight", "1550"))
    cells = {cell.get("id"): cell for cell in graph_root.findall("mxCell")}
    bounds = {cell_id: _mx_geometry(cell) for cell_id, cell in cells.items()}

    svg = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        (
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{page_width}" '
            f'height="{page_height}" viewBox="0 0 {page_width} {page_height}">'
        ),
        "<title>Mentor Market UML Class Diagram</title>",
        "<desc>Rendered directly from Mentor_Market_Class_Diagram.drawio</desc>",
        "<defs>",
        (
            '<marker id="hollowTriangle" markerWidth="16" markerHeight="16" '
            'refX="14" refY="8" orient="auto" markerUnits="userSpaceOnUse">'
            '<path d="M 0 1 L 14 8 L 0 15 Z" fill="#ffffff" stroke="#000000" '
            'stroke-width="1.5"/></marker>'
        ),
        (
            '<marker id="openArrow" markerWidth="14" markerHeight="14" '
            'refX="13" refY="7" orient="auto" markerUnits="userSpaceOnUse">'
            '<path d="M 1 1 L 13 7 L 1 13" fill="none" stroke="#000000" '
            'stroke-width="1.5"/></marker>'
        ),
        (
            '<marker id="hollowDiamond" markerWidth="22" markerHeight="16" '
            'refX="1" refY="8" orient="auto-start-reverse" markerUnits="userSpaceOnUse">'
            '<path d="M 1 8 L 10 1 L 20 8 L 10 15 Z" fill="#ffffff" '
            'stroke="#000000" stroke-width="1.5"/></marker>'
        ),
        (
            '<marker id="filledDiamond" markerWidth="22" markerHeight="16" '
            'refX="1" refY="8" orient="auto-start-reverse" markerUnits="userSpaceOnUse">'
            '<path d="M 1 8 L 10 1 L 20 8 L 10 15 Z" fill="#000000" '
            'stroke="#000000" stroke-width="1.5"/></marker>'
        ),
        "</defs>",
        f'<rect x="0" y="0" width="{page_width}" height="{page_height}" fill="#ffffff"/>',
    ]

    # Draw class boxes first. Their labels and compartment structure come
    # from the HTML stored in each vertex of the Draw.io file.
    text_vertices = []
    for cell in cells.values():
        if cell.get("vertex") != "1":
            continue
        style = _mx_style(cell.get("style"))
        if style.get("opacity") == "0":
            continue
        cell_id = cell.get("id", "")
        x, y, width, height = bounds[cell_id]
        if "text" in style:
            text_vertices.append(cell)
            continue

        parts = _drawio_text_parts(cell.get("value"))
        if not parts:
            continue
        header = parts[0]
        abstract = bool(header and header[0] == "«abstract»")
        class_name = header[-1]
        remaining = parts[1:]
        attributes = []
        operations = []
        if len(remaining) >= 2:
            attributes, operations = remaining[0], remaining[1]
        elif remaining:
            if all(line.startswith("+") for line in remaining[0]):
                operations = remaining[0]
            else:
                attributes = remaining[0]

        svg.append(
            f'<rect x="{x:.2f}" y="{y:.2f}" width="{width:.2f}" height="{height:.2f}" '
            'fill="#ffffff" stroke="#000000" stroke-width="1.5"/>'
        )
        header_height = 54 if abstract else 38
        separator_y = y + header_height
        svg.append(
            f'<line x1="{x:.2f}" y1="{separator_y:.2f}" x2="{x + width:.2f}" '
            f'y2="{separator_y:.2f}" stroke="#000000" stroke-width="1.2"/>'
        )
        if abstract:
            svg.append(_svg_text(x + width / 2, y + 17, "«abstract»", 12, "middle", style="italic"))
            svg.append(
                _svg_text(
                    x + width / 2,
                    y + 40,
                    class_name,
                    14,
                    "middle",
                    "bold",
                    "italic",
                )
            )
        else:
            svg.append(_svg_text(x + width / 2, y + 25, class_name, 14, "middle", "bold"))

        cursor = separator_y + 19
        line_height = 18
        for line in attributes:
            svg.append(_svg_text(x + 9, cursor, line, 12))
            cursor += line_height
        if attributes and operations:
            operations_separator = cursor + 2
            svg.append(
                f'<line x1="{x:.2f}" y1="{operations_separator:.2f}" '
                f'x2="{x + width:.2f}" y2="{operations_separator:.2f}" '
                'stroke="#000000" stroke-width="1.2"/>'
            )
            cursor = operations_separator + 19
        for line in operations:
            svg.append(_svg_text(x + 9, cursor, line, 12))
            cursor += line_height

    # Draw relationships from the endpoints and waypoints stored in the
    # Draw.io file. Labels and multiplicities are retained from the source.
    for cell in cells.values():
        if cell.get("edge") != "1":
            continue
        source_id = cell.get("source")
        target_id = cell.get("target")
        if source_id not in bounds or target_id not in bounds:
            continue
        style = _mx_style(cell.get("style"))
        source_bounds = bounds[source_id]
        target_bounds = bounds[target_id]
        source_center = (
            source_bounds[0] + source_bounds[2] / 2,
            source_bounds[1] + source_bounds[3] / 2,
        )
        target_center = (
            target_bounds[0] + target_bounds[2] / 2,
            target_bounds[1] + target_bounds[3] / 2,
        )
        start = _cell_anchor(source_bounds, style, "exit", target_center)
        end = _cell_anchor(target_bounds, style, "entry", source_center)

        geometry = cell.find("mxGeometry")
        waypoints = []
        if geometry is not None:
            points = geometry.find("Array[@as='points']")
            if points is not None:
                for point in points.findall("mxPoint"):
                    waypoints.append((float(point.get("x", "0")), float(point.get("y", "0"))))
        points = _orthogonal_points(start, end, waypoints)
        point_text = " ".join(f"{x:.2f},{y:.2f}" for x, y in points)
        attributes = [
            'fill="none"',
            'stroke="#000000"',
            'stroke-width="1.5"',
            'stroke-linejoin="miter"',
        ]
        if style.get("dashed") == "1":
            attributes.append('stroke-dasharray="8 8"')
        if style.get("endArrow") == "block" and style.get("endFill") == "0":
            attributes.append('marker-end="url(#hollowTriangle)"')
        elif style.get("endArrow") == "open":
            attributes.append('marker-end="url(#openArrow)"')
        if style.get("startArrow") == "diamondThin":
            marker = "filledDiamond" if style.get("startFill") == "1" else "hollowDiamond"
            attributes.append(f'marker-start="url(#{marker})"')
        svg.append(f'<polyline points="{point_text}" {" ".join(attributes)}/>')

        label = html.unescape(cell.get("value", "")).strip()
        if label:
            label_x, label_y, vertical = _edge_label_position(points)
            label_width = max(34, len(label) * 6.2 + 12)
            label_x += 10 if vertical else 0
            svg.append(
                f'<rect x="{label_x - label_width / 2:.2f}" y="{label_y - 12:.2f}" '
                f'width="{label_width:.2f}" height="18" fill="#ffffff"/>'
            )
            svg.append(_svg_text(label_x, label_y + 2, label, 11, "middle"))

    # Title and multiplicity labels are ordinary text vertices in Draw.io.
    for cell in text_vertices:
        style = _mx_style(cell.get("style"))
        x, y, width, height = bounds[cell.get("id")]
        parts = _drawio_text_parts(cell.get("value"))
        lines = [line for part in parts for line in part]
        if not lines:
            continue
        text = " ".join(lines)
        size = float(style.get("fontSize", "11"))
        weight = "bold" if "<b>" in (cell.get("value") or "") else "normal"
        svg.append(
            _svg_text(
                x + width / 2,
                y + height / 2 + size * 0.35,
                text,
                size,
                "middle",
                weight,
            )
        )

    svg.append("</svg>")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    DRAWIO_SVG.write_text("\n".join(svg), encoding="utf-8")
    subprocess.run(
        [
            "rsvg-convert",
            "--width",
            str(int(page_width * 2)),
            "--height",
            str(int(page_height * 2)),
            "--output",
            str(DOMAIN_DIAGRAM),
            str(DRAWIO_SVG),
        ],
        check=True,
    )


def add_diagram(doc, path, title):
    if title:
        add_heading(doc, title, first=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(9.8))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title or "UML Class Diagram")
    doc_pr.set(
        "descr",
        "Mentor Market UML class diagram rendered from the editable Mentor_Market_Class_Diagram.drawio source.",
    )


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    export_drawio_class_diagram()

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "470 A1 — Mentor Market"
    doc.core_properties.subject = "Functional requirements, class diagram, and architecture analysis"
    doc.core_properties.author = "Mentor Market Project Team"
    doc.core_properties.keywords = "CSE470, Mentor Market, requirements, MVC, class diagram"

    add_heading(doc, "Q1 — Functional and Non-Functional Requirements", first=True)
    add_body(doc, Q1_INTRO)
    add_heading(doc, "Functional Requirements")
    add_heading(doc, "Functional User Requirements")
    add_body(doc, FUNCTIONAL_USER_INTRO)
    for requirement in FUNCTIONAL_USER_REQUIREMENTS:
        add_bullet(doc, requirement)

    add_heading(doc, "Functional System Requirements")
    add_body(doc, FUNCTIONAL_SYSTEM_INTRO)
    for title, requirements in FUNCTIONAL_SYSTEM_REQUIREMENTS:
        add_heading(doc, title)
        for requirement in requirements:
            add_bullet(doc, requirement)

    add_heading(doc, "Non-Functional Requirements")
    add_body(doc, NON_FUNCTIONAL_INTRO)
    for title, requirements in NON_FUNCTIONAL_REQUIREMENTS:
        add_heading(doc, title)
        for requirement in requirements:
            add_bullet(doc, requirement)

    q2_intro = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(q2_intro)
    q2_intro.header.paragraphs[0].clear()
    q2_intro.footer.paragraphs[0].clear()
    add_heading(doc, "Q2 — UML Class Diagram", first=True)
    for explanation in Q2_EXPLANATIONS:
        add_body(doc, explanation)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(landscape, landscape=True, diagram=True)
    landscape.header.paragraphs[0].clear()
    landscape.footer.paragraphs[0].clear()
    add_diagram(doc, DOMAIN_DIAGRAM, None)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(portrait)
    portrait.header.paragraphs[0].clear()
    portrait.footer.paragraphs[0].clear()
    add_heading(doc, "Q3 — Why MVC Is the Best Fit for Mentor Market", first=True)
    add_body(doc, Q3_INTRO)
    add_heading(doc, "Model, View and Controller in Mentor Market")
    for item in Q3_MVC_COMPONENTS:
        add_body(doc, item)

    add_heading(doc, "Routes and the Request Cycle")
    add_body(doc, Q3_ROUTES)
    add_heading(doc, "Example Request Cycle — Publishing a Tutor Service")
    for number, item in enumerate(Q3_REQUEST_CYCLE, start=1):
        add_body(doc, f"{number}. {item}")

    add_heading(doc, "Why MVC Fits Mentor Market")
    for item in Q3_MVC_REASONS:
        add_body(doc, item)

    add_heading(
        doc,
        "A Practical Alternative — Layered Architecture",
    )
    for item in Q3_ALTERNATIVE:
        add_body(doc, item)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
